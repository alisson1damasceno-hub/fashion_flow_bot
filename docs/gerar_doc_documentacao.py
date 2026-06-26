"""
Gera documentacao_projeto.docx — a documentação completa do Fashion Flow Bot,
com as decisões de design e o porquê de cada uma.

Execute: python docs/gerar_doc_documentacao.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

OUT = Path(__file__).parent / "documentacao_projeto.docx"

ROXO = RGBColor(0x76, 0x4A, 0xB0)
VERDE = RGBColor(0x27, 0xAE, 0x60)
LARANJA = RGBColor(0xE6, 0x7E, 0x22)
VERMELHO = RGBColor(0xC0, 0x39, 0x2B)
CINZA = RGBColor(0x55, 0x55, 0x55)


def add_run(par, text, bold=False, italic=False, color=None, mono=False, size=None):
    r = par.add_run(text)
    if bold: r.bold = True
    if italic: r.italic = True
    if color: r.font.color.rgb = color
    r.font.name = "Consolas" if mono else "Calibri"
    if mono and size is None: size = 9.5
    if size: r.font.size = Pt(size)
    return r


def h(doc, text, level=1, color=None):
    par = doc.add_heading(level=level)
    r = par.add_run(text)
    if color: r.font.color.rgb = color
    return par


def p(doc, text=""):
    return doc.add_paragraph(text)


def code(doc, lines):
    par = doc.add_paragraph()
    par.paragraph_format.left_indent = Cm(0.4)
    par.paragraph_format.space_before = Pt(2)
    par.paragraph_format.space_after = Pt(2)
    add_run(par, lines if isinstance(lines, str) else "\n".join(lines),
            mono=True, color=CINZA, size=9)
    pPr = par._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'F4F4F4')
    pPr.append(shd)


def tabela(doc, headers, linhas):
    tab = doc.add_table(rows=len(linhas) + 1, cols=len(headers))
    tab.style = "Light Grid Accent 1"
    for j, lbl in enumerate(headers):
        tab.rows[0].cells[j].text = ""
        add_run(tab.rows[0].cells[j].paragraphs[0], lbl, bold=True)
    for i, row in enumerate(linhas, 1):
        for j, val in enumerate(row):
            tab.rows[i].cells[j].text = str(val)
    return tab


def decisao(doc, codigo, titulo, porque):
    par = p(doc)
    add_run(par, f"{codigo} — {titulo}", bold=True, color=ROXO)
    pq = p(doc)
    add_run(pq, "Por quê: ", bold=True)
    pq.add_run(porque)


# ============================================================
doc = Document()
for s in doc.sections:
    s.left_margin = Cm(2.0); s.right_margin = Cm(2.0)
    s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)

# CAPA
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(t, "\n\n\nFashion Flow Bot\n", bold=True, size=28, color=ROXO)
add_run(t, "Documentação do Projeto\n", bold=True, size=18)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(sub, "\nO que é, como funciona e as decisões de design (com o porquê)\n",
        italic=True, color=CINZA)
add_run(sub, "Equipe: Alisson Damasceno · Bernardo Mota · Márcia Beatriz Costa · Vinícius (oreddd)\n",
        color=CINZA, size=10)
add_run(sub, "2026-06-26\n", color=CINZA, size=10)

p(doc)
caixa = doc.add_paragraph()
caixa.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(caixa, "Chatbot do setor de produção • base em CSV • 147 intenções • CRUD com persistência real • 42 testes",
        bold=True, size=10)

doc.add_page_break()

# 1. VISÃO GERAL
h(doc, "1. Visão geral", level=1, color=ROXO)
p(doc,
  "O Fashion Flow Bot é um chatbot de atendimento do setor de PRODUÇÃO de uma confecção "
  "fictícia. Ele responde dúvidas sobre produção (tecidos, processos, qualidade, "
  "personalização, prazos, cuidados, catálogo, sustentabilidade) e gerencia pedidos "
  "(criar, consultar, alterar, cancelar), tudo pela conversa.")
for item in [
    "Sem IA externa: o cérebro é uma matriz de conhecimento em CSV + regras em Python. É transparente e fácil de manter.",
    "Persistência real: as operações de pedido alteram os arquivos CSV de verdade.",
    "Memória da conversa: o bot lembra do contexto, do nome do cliente e do estado do diálogo (tema 'A Memória do Bot', Semana 3).",
]:
    doc.add_paragraph(item, style="List Bullet")
restr = p(doc)
add_run(restr, "Restrição que guia tudo: ", bold=True)
restr.add_run("somos apenas o bot do setor de produção. Isso decide o que o bot responde "
              "(fabricação) e o que ele encaminha para outro setor (entrega → logística; "
              "preço/pagamento → vendas; devolução → devoluções).")

# 2. COMO RODAR
h(doc, "2. Como rodar", level=1, color=ROXO)
p(doc, "Instale as dependências com pip install -r requirements.txt. Depois:")
tabela(doc, ["Quero...", "Comando"], [
    ["Conversar no terminal", "python main.py"],
    ["Subir a interface web", "python -m uvicorn app:app --reload  →  http://localhost:8000"],
    ["Ver o CRUD funcionando (demo)", "python demo_crud.py"],
    ["Gerenciar a base de conhecimento (dev)", "python gerenciar_intencoes.py"],
    ["Rodar os testes", "python -m unittest test_crud -v"],
])
p(doc, "No terminal, '/contexto' mostra a memória da sessão e 'sair' encerra.")

# 3. ARQUITETURA
h(doc, "3. Arquitetura", level=1, color=ROXO)
h(doc, "3.1 O pipeline (o caminho de uma mensagem)", level=2)
p(doc, "Cada mensagem do usuário passa por esta sequência (em main.py e app.py):")
code(doc, [
    "mensagem do usuário",
    "  -> verificar_seguranca()       bloqueia senha/CVV/cartao antes de tudo   (seguranca.py)",
    "  -> tratar_nome()               no inicio, pergunta e guarda o nome       (cliente.py)",
    "  -> is_despedida / is_casual    atalhos ('tchau', 'ok')                   (contexto.py)",
    "  -> extrair_slots()             acha produto, cor, numero_pedido...       (extractor.py)",
    "  -> merge_com_contexto()        junta com a memoria da conversa           (contexto.py)",
    "  -> classificar()               decide a intencao (usa os PESOS)          (classifier.py)",
    "  -> responder()                 monta a resposta; no CRUD chama pedidos/  (responder.py)",
    "        -> bot/pedidos/<op>      regra de negocio do CRUD",
    "              -> persistencia     le/grava o pedidos.csv",
    "  -> personalizar()              coloca o nome na frente ('Maria, ...')    (cliente.py)",
    "  -> atualizar_sessao_pos_turno() atualiza foco, historico e estados       (contexto.py)",
])
h(doc, "3.2 Mapa de módulos", level=2)
tabela(doc, ["Módulo", "Responsabilidade"], [
    ["bot/loader.py", "Carrega todos os CSVs em memória (um DataFrame por tabela)."],
    ["bot/normalizar.py", "Função única de normalização (minúsculo + sem acento)."],
    ["bot/seguranca.py", "Filtro que bloqueia dados sensíveis (senha, CVV, cartão)."],
    ["bot/extractor.py", "Extrai 'slots' (dados) da mensagem — produto, cor, quantidade, número do pedido..."],
    ["bot/contexto.py", "A memória da conversa: foco atual, histórico, estados, despedida/casual."],
    ["bot/classifier.py", "Decide a intenção por prioridade de regras + peso das intenções."],
    ["bot/responder.py", "Monta a resposta a partir da intenção + dados."],
    ["bot/cliente.py", "Nome do cliente: captura no início e personaliza as respostas."],
    ["bot/estados.py", "O 'mapa de estados': onde a conversa está + o que o usuário quer."],
    ["bot/pedidos/", "O CRUD de pedidos (um arquivo por operação) — ver seção 5.8."],
])
h(doc, "3.3 Duas interfaces, mesmo cérebro", level=2)
p(doc, "main.py é o terminal (loop while, testes rápidos). app.py é a API REST (FastAPI) "
       "que serve o index.html, com uma sessão isolada por aba. As duas chamam o MESMO "
       "pipeline, então qualquer melhoria vale para ambas.")

# 4. OS DADOS
h(doc, "4. Os dados (CSVs)", level=1, color=ROXO)
p(doc, "Tudo que o bot 'sabe' e 'lembra' vive em data/. São dois tipos:")
h(doc, "4.1 Base de conhecimento e slots", level=2)
doc.add_paragraph("intencoes.csv — a matriz de conhecimento (147 intenções). Colunas: "
                  "id_intencao, palavras_chave, tipo_resposta, resposta_padrao, "
                  "pergunta_followup, peso. Cada linha é uma 'regra'.", style="List Bullet")
doc.add_paragraph("slots.csv — documenta os dados extraíveis (produto, cor, quantidade, "
                  "numero_pedido no formato FF-AAAA-NNNN...).", style="List Bullet")
h(doc, "4.2 Tabelas de referência (lookup)", level=2)
p(doc, "Respostas dinâmicas (prazo, preço, compatibilidade...) saem de tabelas lookup_*: "
       "prazo, preço, cor_tecido, compat_tecido_produto, compat_tecido_personalizacao, "
       "tamanho_produto, gramatura, consumo_tecido, capacidade_produtiva, estoque_materiais, "
       "e lookup_etapas.csv (as 6 etapas de fabricação + se cada uma permite alteração).")
h(doc, "4.3 A tabela de pedidos", level=2)
p(doc, "pedidos.csv — os pedidos reais. Colunas: numero_pedido, data_criacao, cliente, "
       "produto, quantidade, cor, tamanho, tecido, personalizacao, etapa_atual, status, "
       "data_prevista, observacao.")
for item in [
    "etapa_atual aponta para uma etapa de lookup_etapas.csv (é o status de fabricação).",
    "status = ciclo macro: em_producao / concluido / cancelado / pausado.",
    "cliente = o dono do pedido (usado na trava de dono).",
]:
    doc.add_paragraph(item, style="List Bullet")

# 5. FUNCIONALIDADES
doc.add_page_break()
h(doc, "5. Funcionalidades em detalhe", level=1, color=ROXO)

h(doc, "5.1 Classificação por intenção + pesos", level=2)
p(doc, "O classifier decide a intenção em camadas de prioridade: regras por verbo (ex.: "
       "'cancelar'), por slot (um número de pedido), palavra-chave do CSV e, por fim, "
       "similaridade (rapidfuzz). Quando várias intenções batem, vence a de MAIOR PESO.")
h(doc, "5.2 Extração de slots", level=2)
p(doc, "O extractor é uma função pura (não mexe na sessão): lê a mensagem e devolve os "
       "dados. Trata negação ('sem bordado'), plurais, e extrai o número do pedido antes da "
       "quantidade (pra 'FF-2024-0123' não virar 'quantidade 2024').")
h(doc, "5.3 Memória da conversa", level=2)
p(doc, "O contexto guarda dois níveis: foco_atual (o assunto de agora) e historico_turnos "
       "(tudo que aconteceu). Quando o produto muda, os slots-filhos (cor, tecido...) do "
       "produto antigo são esquecidos do foco — mas o histórico preserva tudo.")
h(doc, "5.4 Mapa de estados", level=2)
p(doc, "O estados.py traduz a conversa em dois rótulos, atualizados a cada turno: "
       "estado_conversa (OCIOSO, AGUARDANDO_NOME, EM_ASSUNTO, AGUARDANDO_OPCAO, "
       "AGUARDANDO_ID, COLETANDO_PEDIDO) e objetivo_usuario (CONVERSA_SOCIAL, TIRAR_DUVIDA, "
       "SIMULAR_PEDIDO, GERIR_PEDIDO, IR_OUTRO_SETOR, INDEFINIDO).")
h(doc, "5.5 Personalização", level=2)
p(doc, "O cliente.py pergunta o nome no começo da conversa, guarda em nome_cliente, e "
       "prefixa o nome nas respostas ('Maria, ...'). Limpa introduções ('me chamo...') e "
       "usa só o primeiro nome.")
h(doc, "5.6 Filtro de segurança", level=2)
p(doc, "O seguranca.py roda antes de tudo: se a mensagem tem senha, CVV ou número de "
       "cartão, o bot bloqueia o turno e orienta a nunca compartilhar esses dados no chat.")
h(doc, "5.7 Fronteiras de setor", level=2)
p(doc, "O bot responde só o que é de produção. Para o resto, encaminha (intenções setor_*): "
       "preço/pagamento → vendas; entrega/frete → logística; devolução/defeito → devoluções; "
       "matéria-prima → almoxarifado/compras.")
h(doc, "5.8 CRUD de pedidos (bot/pedidos/)", level=2)
p(doc, "Um arquivo por operação, todos apoiados num único persistencia.py:")
tabela(doc, ["Operação", "Arquivo", "O que faz"], [
    ["CREATE", "criar.py", "Coleta os dados na conversa, gera o ID e grava a linha (nasce em modelagem/em_producao)."],
    ["READ", "consultar.py", "Acha o pedido pelo ID e mostra a etapa atual (ou 'não encontrei' = erro 404)."],
    ["UPDATE", "atualizar.py", "alterar_campo (só na modelagem) e avancar_etapa (operador)."],
    ["DELETE", "cancelar.py", "Soft delete: marca status=cancelado, não apaga a linha."],
    ["(infra)", "persistencia.py", "Único que lê/grava o pedidos.csv + e_dono() (trava de dono)."],
])
h(doc, "5.9 Trava de dono (ownership)", level=2)
p(doc, "Como o chat é do CLIENTE, ele só mexe nos pedidos que estão no nome dele. O nome "
       "(capturado pela personalização) é comparado com a coluna cliente do pedido via "
       "persistencia.e_dono() (sem ligar pra maiúscula/acento). Não bateu → 'esse pedido não "
       "está no seu nome'. O avancar_etapa (operador) não tem essa trava.")
h(doc, "5.10 CRUD da base de conhecimento (dev)", level=2)
p(doc, "O gerenciar_intencoes.py (na raiz) é uma ferramenta DE DESENVOLVEDOR para "
       "criar/consultar/atualizar/remover intenções do intencoes.csv (com o peso). NÃO faz "
       "parte do chat — o bot só lê as intenções; quem gerencia são os devs.")

# 6. DECISÕES (a seção principal)
doc.add_page_break()
h(doc, "6. Decisões de design (e o porquê)", level=1, color=ROXO)
p(doc, "Esta é a seção principal: cada decisão importante, com o motivo.")

decisoes = [
    ("D1", "Base de conhecimento em CSV (matriz), sem IA externa.",
     "é o que o material pede (Semanas 1–4); é transparente e fácil de manter — adicionar "
     "uma linha no CSV já ensina uma resposta nova, sem mexer no código."),
    ("D2", "Pipeline modular (um módulo por responsabilidade).",
     "clareza e testabilidade. O projeto é didático e precisa ser entendido e explicado, "
     "então separamos extrair / classificar / responder / contexto / segurança."),
    ("D3", "Pesos nas intenções (coluna peso).",
     "o professor pediu pesos em sala. Servem de desempate: quando várias intenções batem, "
     "a de maior peso vence (a específica ganha da genérica). Antes era 'a primeira do arquivo'."),
    ("D4", "Memória em dois níveis (foco + histórico) com invalidação.",
     "tema da Semana 3 ('a memória do bot'). Resolve o bug de 'slots grudando' — quando o "
     "cliente troca de produto, o bot não pode continuar com a cor/tecido do produto anterior. "
     "O histórico guarda tudo; o foco só o assunto de agora."),
    ("D5", "Mapa de estados explícito (estado_conversa + objetivo_usuario).",
     "o professor falou em 'mapear estados da conversa e do que o usuário quer'. Dar nome aos "
     "estados deixa o contexto inspecionável (/contexto) e demonstrável."),
    ("D6", "SIMULAR_PEDIDO em vez de 'FAZER_ORCAMENTO'.",
     "'orçamento' fechado (com preço e pagamento) é de vendas, não de produção. A produção só "
     "estima (prazo, viabilidade, consumo, preço indicativo) — o nome reflete isso e mantém a "
     "coerência de setor."),
    ("D7", "AGUARDANDO_NOME como estado.",
     "depois de adicionar a captura de nome, faltava representar essa fase; sem ela, o bot "
     "aparecia como OCIOSO enquanto perguntava o nome."),
    ("D8", "Personalização desde o início.",
     "o professor quer interações personalizadas. Encaixa no tema 'memória': o bot guarda o "
     "nome e chama o cliente por ele. Perguntamos no começo para personalizar tudo."),
    ("D9", "CRUD de pedidos conversacional e com persistência real.",
     "requisito do professor — o CRUD tem que acontecer ao longo da conversa, com as "
     "requisições do cliente alterando o CSV de verdade."),
    ("D10", "Duas tabelas: lookup_etapas (referência) + pedidos (dados).",
     "o lookup_pedidos.csv original era, na verdade, a tabela de etapas (não de pedidos). "
     "Separar referência de dado é o padrão limpo; o pedido aponta para a etapa."),
    ("D11", "Um arquivo por operação de CRUD + persistência única (padrão repositório).",
     "clareza > DRY num projeto didático. Cada operação fica fácil de ler e explicar; e como "
     "só o persistencia.py toca o arquivo, trocar CSV por banco no futuro muda um arquivo só."),
    ("D12", "ID sequencial FF-AAAA-NNNN (em vez de aleatório).",
     "também é único, mas fica legível e ordenado por ano. (O material usa GERAR_ID_ALEATORIO; "
     "ambos resolvem o 'gerar ID'.)"),
    ("D13", "Soft delete (cancelar = mudar status, não apagar).",
     "preserva o histórico — mais realista numa fábrica e mais seguro (nada some por acidente). "
     "Tecnicamente, é um UPDATE de status."),
    ("D14", "Alteração travada por etapa (só na modelagem).",
     "espelha a fábrica real — depois do corte, mudar gera retrabalho e custo. A regra mora na "
     "coluna pode_alterar da tabela de etapas."),
    ("D15", "'Status do pedido' = etapa de fabricação.",
     "coerência de setor. Produção sabe em que etapa a peça está; entrega é com a logística (o "
     "bot avisa isso)."),
    ("D16", "'Avançar etapa' é do operador, fora do chat do cliente.",
     "o chat é do cliente, e empurrar a peça na esteira é ação interna da produção. Como painel "
     "de operador ficou fora do escopo, a função existe e é demonstrada por testes/demo, mas o "
     "cliente não dispara pelo chat."),
    ("D17", "CRUD da base de conhecimento é dev-side.",
     "deixar o usuário do bot criar/alterar/apagar intenções seria furo de segurança. O "
     "professor usou a técnica dizendo que os devs gerenciam as intenções — por isso virou uma "
     "ferramenta separada (gerenciar_intencoes.py)."),
    ("D18", "Trava de dono por nome (em vez de login).",
     "login e tela de autenticação ficaram fora do escopo acadêmico. Reaproveitamos o nome já "
     "capturado: cada pedido tem dono e o cliente só vê/mexe nos seus. (Limitação: nome não é "
     "prova de identidade — ver seção 8.)"),
    ("D19", "Filtro de segurança antes de tudo.",
     "dados sensíveis (senha, CVV, cartão) nunca devem ser processados; o bot bloqueia e "
     "orienta, independente do que mais a mensagem tenha."),
    ("D20", "Função normalizar única.",
     "estava copiada em 5 arquivos. Centralizar garante que todos comparem texto do mesmo jeito "
     "(e evita divergência)."),
]
for codigo, titulo, porque in decisoes:
    decisao(doc, codigo, titulo, porque)

# 7. COMO TESTAR
doc.add_page_break()
h(doc, "7. Como testar", level=1, color=ROXO)
doc.add_paragraph("Testes automatizados: python -m unittest test_crud -v — 42 testes "
                  "cobrindo o CRUD de pedidos (com casos de erro e borda), a trava de dono e o "
                  "CRUD de intenções. Rodam em cópias temporárias dos CSVs, então não tocam nos "
                  "dados reais.", style="List Bullet")
doc.add_paragraph("Demonstração: python demo_crud.py — roda uma conversa que exercita os 4 "
                  "CRUDs + a trava de dono (um cliente tentando ver o pedido de outro) e mostra "
                  "o pedidos.csv antes e depois. Reseta a semente no início (repetível).",
                  style="List Bullet")

# 8. LIMITAÇÕES
h(doc, "8. Limitações conhecidas", level=1, color=ROXO)
for item in [
    "Sem autenticação real: a trava de dono usa o nome informado, que não é prova de identidade. Num sistema real teria login. Ficou como limitação consciente (fora do escopo acadêmico).",
    "Preço é indicativo: a produção dá estimativa; preço e pagamento fecham com vendas.",
    "data_prevista usa prazo fixo (20 dias) no CREATE; poderia puxar do lookup_prazo.",
    "NLU simples (palavra-chave + fuzzy): erros de digitação muito fortes podem cair no fallback.",
]:
    doc.add_paragraph(item, style="List Bullet")

# 9. MAPA DE ARQUIVOS
h(doc, "9. Mapa de arquivos", level=1, color=ROXO)
code(doc, [
    "fashion_flow_bot/",
    "├── main.py                  # interface de terminal (loop de conversa)",
    "├── app.py                   # API REST (FastAPI) + serve o index.html",
    "├── index.html / index_alt.html  # frontends web",
    "├── demo_crud.py             # demonstracao do CRUD + trava de dono",
    "├── gerenciar_intencoes.py   # ferramenta DEV: CRUD da base de conhecimento",
    "├── test_crud.py             # 42 testes automatizados (em copias temporarias)",
    "├── requirements.txt",
    "├── bot/",
    "│   ├── loader.py            # carrega os CSVs",
    "│   ├── normalizar.py        # normalizacao de texto (unica)",
    "│   ├── seguranca.py         # filtro de dados sensiveis",
    "│   ├── extractor.py         # extrai slots da mensagem (funcao pura)",
    "│   ├── contexto.py          # memoria da conversa (foco, historico, estados)",
    "│   ├── classifier.py        # decide a intencao (prioridade + pesos)",
    "│   ├── responder.py         # monta a resposta",
    "│   ├── cliente.py           # nome do cliente + personalizacao",
    "│   ├── estados.py           # mapa de estados (conversa + objetivo)",
    "│   └── pedidos/             # CRUD de pedidos",
    "│       ├── persistencia.py  #   le/grava pedidos.csv + e_dono (trava de dono)",
    "│       ├── criar.py         #   CREATE",
    "│       ├── consultar.py     #   READ",
    "│       ├── atualizar.py     #   UPDATE (alterar_campo + avancar_etapa)",
    "│       └── cancelar.py      #   DELETE (soft delete)",
    "├── data/",
    "│   ├── intencoes.csv        # matriz de conhecimento (147 intencoes, com peso)",
    "│   ├── slots.csv            # documentacao dos slots",
    "│   ├── pedidos.csv          # os pedidos reais (com dono)",
    "│   ├── lookup_etapas.csv    # as 6 etapas de fabricacao",
    "│   └── lookup_*.csv         # prazo, preco, compatibilidades, estoque, etc.",
    "└── docs/                    # entregaveis .docx + geradores",
])

# 10. GLOSSÁRIO
h(doc, "10. Glossário", level=1, color=ROXO)
tabela(doc, ["Termo", "O que é"], [
    ["Intenção", "O objetivo por trás da frase (ex.: 'quanto custa?' → preço)."],
    ["Slot", "Um dado extraído da mensagem (ex.: '200 polos' → quantidade=200, produto=polo)."],
    ["Matriz de conhecimento", "A tabela intencoes.csv (intenção × palavras-chave × resposta)."],
    ["Peso", "Prioridade da intenção no desempate (maior peso vence)."],
    ["Foco atual", "Os slots do assunto que está sendo discutido agora."],
    ["Estado da conversa", "Onde o diálogo está (ex.: esperando um menu, coletando um pedido)."],
    ["Objetivo do usuário", "A meta grande por trás da conversa (ex.: tirar dúvida, gerir pedido)."],
    ["CRUD", "Create, Read, Update, Delete — as 4 operações sobre dados."],
    ["Soft delete", "'Apagar' marcando como cancelado, sem remover a linha (preserva histórico)."],
    ["Trava de dono", "Só o dono do pedido (pelo nome) pode consultá-lo/alterá-lo/cancelá-lo."],
    ["Persistência", "Gravar no disco (CSV) para o dado sobreviver depois que o programa fecha."],
])

p(doc)
fim = p(doc)
fim.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(fim, "— fim do documento —", italic=True, color=CINZA)

doc.save(OUT)
print(f"OK: {OUT}")
print(f"Tamanho: {OUT.stat().st_size} bytes")
