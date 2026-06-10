"""
Gera merge_com_equipe.docx — explica o resultado do merge entre nosso
trabalho (Bloco A) e o que a equipe commitou (Bernardo + oreddd).

Execute: python docs/gerar_doc_merge.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

OUT = Path(__file__).parent / "merge_com_equipe.docx"

ROXO = RGBColor(0x76, 0x4A, 0xB0)
VERDE = RGBColor(0x27, 0xAE, 0x60)
LARANJA = RGBColor(0xE6, 0x7E, 0x22)
VERMELHO = RGBColor(0xC0, 0x39, 0x2B)
CINZA = RGBColor(0x55, 0x55, 0x55)


def add_run(p, text, bold=False, italic=False, color=None, mono=False, size=None):
    r = p.add_run(text)
    if bold: r.bold = True
    if italic: r.italic = True
    if color: r.font.color.rgb = color
    r.font.name = "Consolas" if mono else "Calibri"
    if mono and size is None: size = 9.5
    if size: r.font.size = Pt(size)
    return r


def h(doc, text, level=1, color=None):
    p = doc.add_heading(level=level)
    r = p.add_run(text)
    if color: r.font.color.rgb = color
    return p


def p(doc, text=""):
    return doc.add_paragraph(text)


def code(doc, lines):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(0.5)
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    add_run(para, lines if isinstance(lines, str) else "\n".join(lines),
            mono=True, color=CINZA, size=9.5)
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'F4F4F4')
    pPr.append(shd)


# ============================================================
doc = Document()
for s in doc.sections:
    s.left_margin = Cm(2.0); s.right_margin = Cm(2.0)
    s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)

# CAPA
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(t, "\n\n\nFashion Flow Bot\n", bold=True, size=28, color=ROXO)
add_run(t, "Merge com o trabalho da equipe\n", bold=True, size=18)

s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(s, "\nIntegrando nossas correções com o que Bernardo e oreddd commitaram\n",
        italic=True, color=CINZA)
add_run(s, "Bia Costa • 2026-06-10\n", color=CINZA, size=10)

doc.add_page_break()

# 1. Contexto
h(doc, "1. O que aconteceu", level=1)
p(doc,
  "Enquanto a gente fechava as correções localmente, a equipe também trabalhou. "
  "Foram dois trabalhos paralelos diferentes:")
doc.add_paragraph(
    "Bernardo (commit 253e0ff no main): adicionou Procfile e index(1).html. "
    "NÃO mexeu no código do bot.",
    style="List Bullet")
doc.add_paragraph(
    "Oreddd (branch Front, divergente): adicionou histórico de intenções como mecanismo de "
    "desambiguação, expandiu MUITO os menus do responder, criou um index.html alternativo "
    "(visual azul), e um iniciar.bat. Não atacou os bugs críticos.",
    style="List Bullet")
p(doc,
  "Como nossos arquivos modificados e os deles não se sobrepõem brutalmente, dava pra "
  "fazer um merge limpo. Foi isso que fizemos.")

# 2. Mapa do que veio de onde
h(doc, "2. De onde veio cada coisa", level=1)

tab = doc.add_table(rows=11, cols=3)
tab.style = "Light Grid Accent 1"
hdr = ["Componente", "Vem de", "Por quê"]
for j, lbl in enumerate(hdr):
    tab.rows[0].cells[j].text = ""
    add_run(tab.rows[0].cells[j].paragraphs[0], lbl, bold=True)

linhas = [
    ("bot/contexto.py", "Nosso", "Tem foco_atual + historico_turnos + invalidação por dependência. Solo do oreddd não resolveria CRÍTICO 1."),
    ("bot/extractor.py", "Nosso", "Função pura, negação, plurais, fix do bug do numero_pedido."),
    ("bot/classifier.py", "Nosso", "Regras por verbo primeiro, slots_turno vs efetivos, herança via ultimo_assunto."),
    ("bot/responder.py", "Nosso + menus do oreddd", "Nossa lógica de resposta + dicionário de menus expandido (qualidade, personalização aninhada, sustentabilidade, manutenção, produção, catálogo, sugestão, tecidos, prazos, etapas)."),
    ("main.py", "Nosso", "Pipeline em 5 etapas + comando /contexto pra debug."),
    ("app.py", "Nosso (com adaptação do oreddd)", "Pipeline novo + endpoint /sessao/{id} + serve index.html na raiz se existir."),
    ("front_gradio.py", "Nosso", "gr.State() por usuário + painel de memória."),
    ("index.html", "Bernardo (commit 253e0ff)", "Frontend HTML standalone, tema escuro com paleta verde Spotify-style, font Inter. JS chama nossa API."),
    ("Procfile", "Bernardo (commit 253e0ff)", "Pra deploy automático no Heroku/Render."),
    ("docs/", "Nosso", "auditoria_bugs.docx, correcoes_aplicadas.docx, merge_com_equipe.docx."),
]
for i, row in enumerate(linhas, 1):
    for j, val in enumerate(row):
        tab.rows[i].cells[j].text = val

# 3. Comparação das 3 versões
doc.add_page_break()
h(doc, "3. Comparação rápida das 3 versões", level=1)

p(doc, "Nas linhas abaixo: ✅ = consertado/feito, ❌ = problema continua, — = não aplicável.")

t2 = doc.add_table(rows=11, cols=4)
t2.style = "Light Grid Accent 1"
hdr = ["Tópico", "Nosso (final)", "Bernardo (main)", "oreddd (Front)"]
for j, lbl in enumerate(hdr):
    t2.rows[0].cells[j].text = ""
    add_run(t2.rows[0].cells[j].paragraphs[0], lbl, bold=True)
linhas = [
    ("CRÍTICO 1: slots grudam quando muda produto", "✅", "❌", "❌"),
    ("CRÍTICO 2: ano de pedido vira quantidade", "✅", "❌", "❌"),
    ("CRÍTICO 3: menu corrompe quantidade", "✅", "❌", "❌"),
    ("CRÍTICO 4: 'obrigado...' reseta sessão", "✅", "❌", "❌"),
    ("CRÍTICO 5/6: numero_pedido e cancelar", "✅", "❌", "❌"),
    ("CRÍTICO 7: Gradio gr.State", "✅", "❌", "—"),
    ("CRÍTICO 22: negação 'sem bordado'", "✅", "❌", "❌"),
    ("Menus expandidos (15 menus aninhados)", "✅ (do oreddd)", "❌", "✅"),
    ("Frontend HTML standalone", "✅ (do Bernardo)", "✅", "✅"),
    ("Procfile pra deploy", "✅ (do Bernardo)", "✅", "✅"),
]
for i, row in enumerate(linhas, 1):
    for j, val in enumerate(row):
        cell = t2.rows[i].cells[j]
        cell.text = val
        # colore os checkmarks
        if val == "✅":
            for r in cell.paragraphs[0].runs:
                r.font.color.rgb = VERDE
                r.bold = True
        elif val == "❌":
            for r in cell.paragraphs[0].runs:
                r.font.color.rgb = VERMELHO
                r.bold = True

p(doc)
px = p(doc)
add_run(px, "Conclusão: ", bold=True)
px.add_run("o trabalho da equipe foi orientado pra deploy e cobertura de menus (boa coisa!), "
           "mas não atacou os bugs sérios da auditoria. Combinando os dois, temos o melhor "
           "de cada lado.")

# 4. Como o frontend HTML conversa com a API
h(doc, "4. Como o frontend HTML funciona", level=1)

p(doc, "O index.html é um single-page totalmente independente do Gradio. Ele tem um input "
       "de URL da API (default http://localhost:8000) e faz fetch direto:")

code(doc, [
    "// trecho do index.html",
    'const SESSION_ID = "sess_" + Math.random().toString(36).slice(2, 9);',
    "",
    "// pra cada mensagem do usuário:",
    "const res = await fetch(`${api}/chat`, {",
    "    method: 'POST',",
    "    headers: { 'Content-Type': 'application/json' },",
    "    body: JSON.stringify({",
    "        sessao_id: SESSION_ID,",
    "        mensagem: text,",
    "    }),",
    "});",
    "const { resposta } = await res.json();",
])

p(doc, "Cada aba aberta gera um SESSION_ID novo (aleatório). Como o nosso app.py guarda as "
       "sessões num dict por sessao_id, cada aba é uma conversa isolada — equivalente ao "
       "gr.State do Gradio. Bug 7 não acontece no HTML porque cada cliente identifica a si "
       "mesmo.")

# 5. Como rodar
h(doc, "5. Como rodar a versão consolidada", level=1)

px = p(doc)
add_run(px, "Opção 1 — Frontend HTML (recomendado pra demo):", bold=True)
code(doc, [
    "uvicorn app:app --reload",
    "# abre http://localhost:8000  → aparece o index.html",
])

px = p(doc)
add_run(px, "Opção 2 — Gradio (recomendado pra apresentação técnica):", bold=True)
code(doc, ["python front_gradio.py"])

px = p(doc)
add_run(px, "Opção 3 — Terminal puro (debug rápido):", bold=True)
code(doc, ["python main.py    # digite /contexto pra ver a memória"])

px = p(doc)
add_run(px, "Para deploy: ", bold=True)
px.add_run("o Procfile já está pronto. Em Render/Railway/Heroku, basta apontar pro repo e a "
           "plataforma cuida do resto.")

# 6. Próximos passos
h(doc, "6. Próximos passos sugeridos", level=1)
doc.add_paragraph(
    "Commitar tudo num commit só com a mensagem: 'merge: bloco A de correções + frontend "
    "HTML da equipe (Bernardo) + menus expandidos (oreddd)'.",
    style="List Number")
doc.add_paragraph(
    "Abrir o navegador em http://localhost:8000 após subir uvicorn e testar todos os 15 "
    "cenários por clique mesmo.",
    style="List Number")
doc.add_paragraph(
    "Conversar com a equipe sobre fazer merge da branch Front no main, OU se vão preferir "
    "rebatear a branch Front sobre nosso trabalho.",
    style="List Number")
doc.add_paragraph(
    "Atualizar o lookup_estoque_materiais.csv: as datas de reposição estão em 2025.",
    style="List Number")
doc.add_paragraph(
    "Se sobrar tempo: absorver o historico_intencoes do oreddd como complemento ao nosso "
    "historico_turnos — ele tem uma função 'intencoes_relacionadas' que pode ajudar em "
    "fallback de mensagens vagas.",
    style="List Number")

p(doc)
fim = p(doc)
fim.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(fim, "— fim do documento —", italic=True, color=CINZA)

doc.save(OUT)
print(f"OK: {OUT}")
print(f"Tamanho: {OUT.stat().st_size} bytes")
