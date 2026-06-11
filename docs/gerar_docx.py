"""
Gera auditoria_bugs.docx a partir do conteudo do plan file.
Execute: python docs/gerar_docx.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

OUT = Path(__file__).parent / "auditoria_bugs.docx"

# ------------- helpers -------------

COR_CRITICO = RGBColor(0xC0, 0x39, 0x2B)   # vermelho
COR_ALTO    = RGBColor(0xE6, 0x7E, 0x22)   # laranja
COR_MEDIO   = RGBColor(0xF1, 0xC4, 0x0F)   # amarelo escuro
COR_BAIXO   = RGBColor(0x27, 0xAE, 0x60)   # verde
COR_OK      = RGBColor(0x16, 0xA0, 0x85)   # verde-azulado
COR_CINZA   = RGBColor(0x55, 0x55, 0x55)
COR_PRETO   = RGBColor(0x00, 0x00, 0x00)

def add_run(p, text, bold=False, italic=False, color=None, font="Calibri", size=None, mono=False):
    r = p.add_run(text)
    if bold: r.bold = True
    if italic: r.italic = True
    if color: r.font.color.rgb = color
    r.font.name = "Consolas" if mono else font
    if mono and size is None: size = 9.5
    if size: r.font.size = Pt(size)
    return r

def h(doc, text, level=1, color=None):
    p = doc.add_heading(level=level)
    r = p.add_run(text)
    if color: r.font.color.rgb = color
    return p

def p(doc, text=""):
    return doc.add_paragraph(text)

def code_block(doc, lines):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(0.6)
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    add_run(para, lines if isinstance(lines, str) else "\n".join(lines),
            mono=True, color=COR_CINZA, size=9.5)
    # sombreado leve no parágrafo
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'F4F4F4')
    pPr.append(shd)
    return para

def sev_badge(p, severidade):
    cores = {
        "CRITICO": COR_CRITICO, "ALTO": COR_ALTO,
        "MEDIO": COR_MEDIO, "BAIXO": COR_BAIXO,
        "DESIGN": COR_CRITICO,
    }
    add_run(p, f"[{severidade}] ", bold=True, color=cores.get(severidade, COR_PRETO))

def bug(doc, numero, severidade, titulo, repro=None, causa=None, codigo_causa=None,
        por_que=None, como_consertar=None, codigo_extra=None):
    # heading do bug
    head = doc.add_heading(level=3)
    sev_badge(head, severidade)
    head.add_run(f"#{numero} — {titulo}").bold = True

    if repro:
        psub = p(doc)
        add_run(psub, "Como reproduz:", bold=True)
        code_block(doc, repro)

    if causa or codigo_causa:
        psub = p(doc)
        add_run(psub, "Causa-raiz:", bold=True)
        if causa:
            p(doc, causa)
        if codigo_causa:
            code_block(doc, codigo_causa)

    if codigo_extra:
        code_block(doc, codigo_extra)

    if por_que:
        psub = p(doc)
        add_run(psub, "Por que importa: ", bold=True)
        psub.add_run(por_que)

    if como_consertar:
        psub = p(doc)
        add_run(psub, "Como consertar: ", bold=True)
        psub.add_run(como_consertar)

# ------------- documento -------------

doc = Document()

# margens
for s in doc.sections:
    s.left_margin = Cm(2.0); s.right_margin = Cm(2.0)
    s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)

# CAPA
titulo = doc.add_paragraph()
titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(titulo, "\n\n\nFashion Flow Bot\n", bold=True, size=28, color=COR_CRITICO)
add_run(titulo, "Auditoria Completa de Bugs e Lógica\n", bold=True, size=18)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(sub, "\nProjeto acadêmico — Setor de Produção\n", italic=True, size=12, color=COR_CINZA)
add_run(sub, "Bia Costa  •  2026-06-09  •  branch main (commit 9233d6e)\n", size=11, color=COR_CINZA)

p(doc)
caixa = doc.add_paragraph()
caixa.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(caixa, "33 problemas identificados  •  8 críticos  •  11 altos  •  10 médios  •  4 baixos",
        bold=True, size=11)

doc.add_page_break()

# ===== SECAO 0: contexto =====
h(doc, "1. Contexto e objetivo", level=1)
p(doc,
  "O fashion_flow_bot é um chatbot acadêmico do setor de Produção da empresa fictícia Fashion Flow. "
  "A entrega da semana é melhorar a detecção de intenção e salvar o contexto da conversa, "
  "deixando o bot parecido com pessoa real — tema 'Semana 1: O Poder do Contexto' do guia (pág. 14-16, "
  "analogia do garçom).")
p(doc,
  "A equipe já implementou a base de contexto (bot/contexto.py, herança parcial em classifier.py, sessão "
  "por id em app.py). Funciona em muitos casos, mas duas baterias de testes encontraram 33 problemas reais, "
  "alguns graves o suficiente pra fazer o bot dar respostas absurdas.")

obs = p(doc)
add_run(obs, "Nota de método: ", bold=True)
obs.add_run("os testes foram rodados via scripts Python invocando o pipeline ")
add_run(obs, "extrair_slots → classificar → responder", mono=True)
obs.add_run(" direto. A interface (Gradio/main.py) ainda não foi testada interativamente — "
            "recomendação é fazer isso depois de aplicar os fixes do Bloco A.")

# ===== SECAO 1: sumário =====
h(doc, "2. Sumário executivo", level=1)

tab = doc.add_table(rows=7, cols=5)
tab.style = "Light Grid Accent 1"
hdr = tab.rows[0].cells
for i, lbl in enumerate(["Categoria", "Crítico", "Alto", "Médio", "Baixo"]):
    hdr[i].text = ""
    add_run(hdr[i].paragraphs[0], lbl, bold=True)

linhas = [
    ("Contexto / sessão", 3, 3, 2, 1),
    ("Extração de slots", 2, 3, 3, 1),
    ("Classificação de intenção", 2, 3, 2, 0),
    ("Resposta", 0, 2, 2, 1),
    ("Dados (CSV)", 0, 1, 1, 0),
    ("Design / arquitetura", 1, 1, 0, 0),
]
for i, row in enumerate(linhas, start=1):
    for j, val in enumerate(row):
        tab.rows[i].cells[j].text = str(val)

p(doc)
nota = p(doc)
add_run(nota, "Os 8 críticos respondem por ~80% das experiências ruins. Os outros são polimento.",
        italic=True, color=COR_CINZA)

# ===== SECAO 2: mapa de bugs =====
doc.add_page_break()
h(doc, "3. Mapa de bugs", level=1)
p(doc, "Cada bug abaixo tem: número, severidade, sintoma reproduzível, causa-raiz no código, "
       "por que importa e como consertar.")

# --- CRÍTICOS ---
h(doc, "3.1 Críticos", level=2, color=COR_CRITICO)

bug(doc, 1, "CRITICO", "Slots grudam pra sempre na sessão (herança cega)",
    repro=[
        "Você: preço de 100 polos com bordado em algodão",
        "Bot:  Para 100 polos com bordado: R$ 88/peça",
        "Você: agora 200 calças jeans",
        "Bot:  Para 200 PEÇAS DE POLO COM BORDADO: R$ 84/peça  ← ???",
        "Você: quanto custa?",
        "Bot:  Para 200 polo com bordado: R$ 84/peça           ← ainda polo",
    ],
    causa="bot/contexto.py:14-23",
    codigo_causa=[
        "def merge_slots(sessao, slots_novos):",
        "    for chave, valor in slots_novos.items():",
        "        sessao['slots_acumulados'][chave] = valor",
        "    return sessao['slots_acumulados']",
    ],
    por_que="O merge só adiciona/atualiza, nunca remove. O bot fala 'polo com bordado' quando o "
            "cliente já mudou pra calças — exatamente o oposto de 'parecer pessoa real'.",
    como_consertar="Em merge_slots, aplicar invalidação por dependência: quando produto muda, "
                   "zerar personalização/cor/grade (são propriedades do produto antigo). Slots de "
                   "pergunta única (numero_pedido, prazo_desejado, urgente) devem viver 1-2 turnos.")

bug(doc, 2, "CRITICO", "Quantidade extraída de dentro do número de pedido",
    repro=["Você: meu pedido FF-2024-0123 está em qual etapa",
           "Bot:  [slots] quantidade=2024, numero_pedido=FF-2024-0123"],
    causa="bot/extractor.py:20 — o regex deixa a unidade opcional",
    codigo_causa=[r"match = re.search(r'\b(\d+)\s*(pecas?|...)?\b', t)"],
    por_que="Esse quantidade=2024 herda nos próximos turnos. Cliente pergunta 'quanto custa?' "
            "e o bot orça 2024 peças. Em cascata, alimenta os Críticos 3, 5 e 6.",
    como_consertar="Extrair numero_pedido antes e tirar o trecho da string antes de procurar "
                   "quantidade. Ou exigir unidade obrigatória junto do número.")

bug(doc, 3, "CRITICO", "Menu numérico corrompe slot de quantidade",
    repro=["Você: preço de 100 polos              → quantidade=100",
           "Você: qualidade                        → menu mostrado",
           "Você: 1                                → escolheu opção 1",
           "Você: e o prazo?",
           "Bot:  Para 1 PEÇA de polo: 18-22 dias  ← virou 1 peça!"],
    causa="O extractor extrai quantidade=1 do dígito '1' digitado pro menu, sobrescrevendo o 100.",
    por_que="Quebra qualquer fluxo onde o cliente desvia pra um menu auxiliar e volta. Bot soa incoerente.",
    como_consertar="Quando aguardando_opcao está ativo, não extrair slots normais — ou ignorar "
                   "a chave 'quantidade' nesse turno.")

bug(doc, 4, "CRITICO", "is_despedida reseta a sessão em frases longas",
    repro=["is_despedida('obrigado, agora me fala o prazo')    = True   ← RESETA!",
           "is_despedida('valeu pela info, mas e o preço?')    = True   ← RESETA!",
           "is_despedida('tchau preciso saber uma última coisa') = True ← RESETA!"],
    causa="bot/contexto.py:46 — uso de substring match",
    codigo_causa=["return any(p in t for p in palavras)   # substring!"],
    por_que="Qualquer mensagem que contenha 'obrigado'/'valeu'/'tchau' em qualquer posição zera a "
            "sessão. Quebra fluxo natural de fala.",
    como_consertar="Usar match exato (t.strip() == palavra) ou exigir que despedida seja a "
                   "intenção dominante. Tirar 'obrigado/valeu' — esses são casuais, não fim de conversa.")

bug(doc, 5, "CRITICO", "numero_pedido transforma qualquer mensagem em status_pedido",
    repro=["Você: pedido FF-2024-0123          → numero_pedido salvo",
           "Você: tem algodão em estoque?       → bot responde sobre o pedido!"],
    causa="bot/classifier.py:44-45",
    codigo_causa=["if slots.get('numero_pedido'):",
                  "    return 'status_pedido'"],
    por_que="Combinado com Crítico 1, deixa o bot preso num assunto. Cliente pergunta de estoque, "
            "bot insiste no pedido antigo.",
    como_consertar="numero_pedido é slot de turno único. A regra deve checar se foi mencionado "
                   "NESSA mensagem específica, não na sessão.")

bug(doc, 6, "CRITICO", "Cancelamento vira status_pedido",
    repro=["Você: quero CANCELAR meu pedido FF-2025-0001",
           "Bot:  Para consultar o STATUS do pedido FF-2025-0001..."],
    causa="Mesma do Crítico 5. A regra de numero_pedido roda antes de qualquer regra que reconheça "
          "a palavra 'cancelar'.",
    como_consertar="Adicionar regra de verbo: if 'cancelar' in t: return 'cancelar_pedido' "
                   "ANTES da regra do numero_pedido. Ordenar regras por intenção (verbo) antes "
                   "de regras puramente por slot.")

bug(doc, 7, "CRITICO", "Slot global no Gradio compartilha entre usuários",
    causa="front_gradio.py:10-14 — dict no escopo do módulo",
    codigo_causa=["sessao = {",
                  "    'ativa': False,",
                  "    'ultimo_assunto': None,",
                  "    'slots_acumulados': {},",
                  "}"],
    por_que="Todo usuário compartilha a mesma sessão. Em qualquer demo com mais de uma aba aberta "
            "o bug aparece. Em produção seria desastre de privacidade.",
    como_consertar="Usar gr.State() por sessão de chat, ou padrão do app.py (dict de sessões por id).")

bug(doc, 22, "CRITICO", "Negação ignorada: 'sem bordado' extrai personalização=bordado",
    repro=["Você: preço de 100 polos SEM bordado",
           "Bot:  Para 100 polos COM BORDADO: R$ 88/peça"],
    causa="bot/extractor.py:65-67 só procura a palavra 'bordado' — não olha a palavra anterior.",
    como_consertar="Detectar `sem\\s+(bordado|silk|...)` antes da busca normal e marcar "
                   "personalizacao='nenhuma'. Ou olhar 1-2 palavras antes do match.")

# --- ALTOS ---
h(doc, "3.2 Altos", level=2, color=COR_ALTO)

bug(doc, 8, "ALTO", "Frases ambíguas extraem quantidade E prazo do mesmo número",
    repro=["'preciso em 100 dias'          → quantidade=100, prazo_desejado=100",
           "'pedido com prazo de 5 dias'   → quantidade=5,   prazo_desejado=5"],
    causa="Dois regex independentes em extractor.py lendo o mesmo texto sem se comunicar.",
    por_que="Afeta validade do orçamento — bot calcula preço de '5 peças' quando o cliente só "
            "falou 'prazo de 5 dias'.",
    como_consertar="Depois de extrair prazo_desejado ou numero_pedido, descartar quantidade se "
                   "vem do mesmo trecho.")

bug(doc, 9, "ALTO", "Follow-up sem número não herda intenção",
    repro=["Você: preço de 200 polos       → ok",
           "Você: e com bordado?            → bot mostra info GENÉRICA",
           "Você: e em algodão?             → idem"],
    causa="bot/classifier.py:25-32 exige número na mensagem nova pra herdar assunto.",
    como_consertar="Se mensagem adiciona pelo menos um slot novo (personalização, tecido, cor) e "
                   "último assunto era preço/prazo/viabilidade, herdar o assunto.")

bug(doc, 10, "ALTO", "selecao_opcao vira ultimo_assunto e quebra herança seguinte",
    causa="main.py:44 — atualiza ultimo_assunto pra 'selecao_opcao', perdendo o assunto real.",
    como_consertar="Quando intencao == 'selecao_opcao', não atualizar ultimo_assunto. Ou setar "
                   "com a sub-intenção real escolhida.")

bug(doc, 11, "ALTO", "qtd_grande_volume come a viabilidade em pedidos grandes",
    repro=["'consigo 100 peças em 5 dias'   → viabilidade_producao  ✓",
           "'consigo 600 peças em 30 dias'  → qtd_grande_volume     ✗ ignora prazo",
           "'consigo 1000 peças em 10 dias' → qtd_grande_volume     ✗"],
    causa="Quando não tem produto, classifier cai no Passo 2 (keywords CSV) e qtd_grande_volume "
          "pega '500+'/'1000' antes que viabilidade seja considerada.",
    como_consertar="Regra adicional: se há prazo_desejado E quantidade, retornar viabilidade_producao "
                   "(sem exigir produto).")

bug(doc, 12, "ALTO", "tem prazo de 10 dias? duplica número em quantidade e prazo",
    causa="Variante do Alto 8. Mesma causa, mesmo fix.")

bug(doc, 23, "ALTO", "Mistura de tecido entre produtos diferentes",
    repro=["Você: e em algodão pima?           → ok pra camiseta",
           "Você: quanto fica em moletom?      → bot oferece MOLETOM EM ALGODÃO PIMA"],
    por_que="Moletom em algodão pima é absurdo (moletom usa moletom_flanelado/peluciado). Tecido "
            "ficou grudado do contexto da camiseta. Variante do CRÍTICO 1.")

bug(doc, 24, "ALTO", "'Volta pro orçamento de X' não restaura contexto antigo",
    repro=["[contexto: 200 moletom + DTF + algodão pima]",
           "Você: volta pro orçamento de camiseta",
           "Bot:  Para 200 camisetas COM DTF EM ALGODÃO PIMA"],
    por_que="Sem histórico, 'voltar' é impossível. O bot acerta o produto (porque keyword) mas "
            "mantém tudo do contexto novo.",
    como_consertar="Implementar historico_turnos no contexto e detectar palavras de revisitação "
                   "('voltar', 'anterior', 'sobre o X de antes').")

bug(doc, 25, "ALTO", "Múltiplas personalizações: só sobrevive a última",
    repro=["Você: preço de 100 polos com bordado E silkscreen",
           "Bot:  Para 100 polos COM SILKSCREEN: R$ 81/peça   ← sumiu o bordado"],
    causa="extractor.py:65-67 itera e sobrescreve. A última chave matchada vence.",
    como_consertar="Aceitar lista, OU pedir esclarecimento ('você quer bordado E silk juntos, ou "
                   "só uma?'). Recomendo a segunda — mais humano.")

bug(doc, 30, "ALTO", "Produto inexistente vira setor_vendas",
    repro=["'preço de 100 ovellas'    → setor_vendas (genérico)",
           "'preço de 100 peças'       → setor_vendas (não pergunta de qual produto)"],
    como_consertar="Quando há quantidade mas não há produto detectado, devolver pergunta "
                   "clarificadora.")

# --- MÉDIOS ---
h(doc, "3.3 Médios", level=2, color=COR_MEDIO)

bug(doc, 13, "MEDIO", "Plurais 'vestidos' e 'calças' não detectam produto",
    repro=["'quero 30 vestidos'  → produto não detectado",
           "'30 calças'           → produto não detectado"],
    causa="extractor.py:25-49 só tem 'vestido midi', 'vestido longo', 'calca jeans', "
          "'calca alfaiataria' — sem 'vestido' nem 'calca' soltos.",
    como_consertar="Acrescentar entradas genéricas, ou usar regex com s? (a gambiarra do "
                   "'moletons' já mostra que dict hardcoded escala mal).")

bug(doc, 14, "MEDIO", "Substring no extrator de personalização",
    repro=["'almofada silkada' → extrai personalizacao=silkscreen (porque 'silk' é substring)"],
    causa="extractor.py:65-67 usa `if chave in t` (sem word boundary).",
    como_consertar=r"Trocar pra re.search(rf'\b{re.escape(chave)}\b', t).")

bug(doc, 15, "MEDIO", "Slot dinâmico criado pelo menu polui o contexto",
    causa="responder.py:60 cria slot 'qualidade' (e similares) que NUNCA é lido depois.",
    como_consertar="Salvar a escolha em outro campo (ex: ultima_subintencao) ou simplesmente não "
                   "persistir — o responder já retornou a resposta correta.")

bug(doc, 16, "MEDIO", "Letra 'M' como tamanho casa em muitos contextos",
    repro=["'tem em tamanho M?'   → grade=adulto ✓",
           "'consigo produzir m peças?' → grade=adulto ✗"],
    causa=r"extractor.py:125 usa re.search(r'\b(pp|p|m|g|gg|xgg)\b', t).",
    como_consertar="Exigir prefixo 'tamanho' ou 'grade' antes da letra.")

bug(doc, 17, "MEDIO", "Slot 'urgente' é extraído mas nunca usado",
    repro=["'urgente! preciso de 100 camisetas em 5 dias'",
           "Bot:  Para 100 camisetas: 18-22 dias  (resposta genérica, ignora urgência)"],
    como_consertar="Ou usar o slot (mensagem dedicada com info de taxa de urgência 20-40%), ou "
                   "tirar do extractor.")

bug(doc, 18, "MEDIO", "Produto pega o último em vez do mais mencionado",
    repro=["'a regata fica boa com legging?' → produto=legging (ignora regata!)"],
    causa="O for itera no dict, primeiro match vence. Ordem de inserção importa — 'legging' vem "
          "antes de 'regata'.",
    como_consertar="Pegar o produto mais à esquerda na frase (re.search + .start()), ou listar "
                   "candidatos e desambiguar.")

bug(doc, 26, "MEDIO", "'preço de 1 camiseta' responde 'Para 1 peças'",
    causa="Plural fixo nos templates do responder.py.",
    como_consertar="`peça` quando quantidade==1, `peças` caso contrário.")

bug(doc, 27, "MEDIO", "Erro de digitação no produto perde o contexto",
    repro=["'preço de 100 camizetas' → setor_vendas (perdeu produto)",
           "'preço de 100 moletoom'  → sobrevive (fuzzy no CSV)",
           "'preço de 100 moletons'  → sobrevive (gambiarra)"],
    causa="Detecção de produto é exata. Algumas grafias erradas funcionam por sorte.",
    como_consertar="Aplicar rapidfuzz também no extractor de produto com threshold ~85%.")

bug(doc, 28, "MEDIO", "Formato alternativo de número de pedido vira fallback total",
    repro=["'pedido FF-25-0010' (ano com 2 dígitos)  → fallback",
           "'pedido FF2025 0010' (sem hífen)         → fallback",
           "'meu pedido 12345'                       → status sem número"],
    como_consertar="Regex relaxado pra detectar tentativas; responder 'o formato é FF-AAAA-NNNN, "
                   "você quis dizer...?'.")

bug(doc, 29, "MEDIO", "status_pedido nunca usa lookup_pedidos.csv",
    causa="responder.py:97-121 sempre devolve 'fale com logística', mesmo o CSV tendo 6 etapas "
          "com pode_alterar + descrição + observação.",
    como_consertar="Mesmo sem saber em qual etapa o pedido X está, o bot pode listar as etapas "
                   "típicas e dizer onde altera (modelagem) e onde já travou (corte em diante).")

bug(doc, 31, "MEDIO", "Quantidade extrapolada (99999) vira 'consulte vendas'",
    repro=["'preço de 99999 camisetas' → não tem range, mensagem genérica do CSV"],
    como_consertar="Mensagem específica pra quantidades muito altas (acima do qtd_max do CSV).")

bug(doc, 33, "MEDIO", "Erro de digitação sobrevive por causa do contexto",
    repro=["[contexto: 100 camisetas com bordado]",
           "Você: qro saber preço (gíria + abreviação)",
           "Bot:  Para 100 camisetas com bordado: R$ 35/peça  ← acertou por acaso"],
    por_que="O classifier não sabe quando está chutando. Se cliente tivesse mudado de produto, "
            "manteria camiseta sem aviso.",
    como_consertar="Medir confiança da classificação. Quando baixa, perguntar em vez de assumir.")

# --- BAIXOS ---
h(doc, "3.4 Baixos", level=2, color=COR_BAIXO)

bug(doc, 19, "BAIXO", "Dados de estoque com datas no passado",
    causa="lookup_estoque_materiais.csv tem previsao_reposicao=2025-06-02 para jeans e "
          "2025-05-28 para linho. Hoje é 2026-06-09.",
    como_consertar="Atualizar o CSV antes de apresentar.")

bug(doc, 20, "BAIXO", "Inconsistência slots.csv vs código",
    causa="slots.csv usa nomes diferentes do código: tipo_personalizacao (CSV) vs personalizacao "
          "(código); tamanho (CSV) vs grade (código).",
    como_consertar="Atualizar a documentação ou o código pra ficarem alinhados.")

bug(doc, 32, "BAIXO", "Dois números de pedido na mesma frase: só pega o primeiro",
    repro=["'o pedido FF-2025-0001 E o pedido FF-2025-0002 estão em qual etapa?'",
           "Bot:  [só fala do FF-2025-0001]"],
    como_consertar="re.findall em vez de re.search. Responder 'vou consultar os dois' ou "
                   "'um por vez'.")

# --- DESIGN ---
h(doc, "3.5 Design / Arquitetura", level=2, color=COR_CRITICO)

bug(doc, 21, "DESIGN", "extrair_slots tem efeito colateral (modifica a sessão)",
    causa="extractor.py:167-169 chama merge_slots dentro da função de extração.",
    por_que="Função impura. O classifier recebe slots já mesclados sem saber o que é novo vs "
            "antigo. É o problema arquitetural por trás dos CRÍTICOS 1 e 3.",
    como_consertar="Separar em duas funções: extrair_slots() pura (só lê), e "
                   "merge_com_contexto(slots, sessao) que retorna o estado efetivo.")

# ===== SECAO 3: design recomendado =====
doc.add_page_break()
h(doc, "4. Design recomendado para o contexto", level=1)

p(doc,
  "A solução para o conflito 'esquecer vs revisitar' não é zerar slots — é separar foco atual de "
  "memória da conversa.")

code_block(doc, [
    "sessao = {",
    "    'foco_atual': {                ← curto prazo (assunto atual)",
    "        'produto': 'calca_jeans',",
    "        'quantidade': 200,",
    "        'tecido': 'jeans',",
    "    },",
    "    'historico_turnos': [          ← longo prazo (memória completa)",
    "        {",
    "            'msg_original': 'preço de 100 polos com bordado em algodão',",
    "            'intencao': 'combinado_preco_qtd_produto',",
    "            'slots': {'produto':'polo', 'quantidade':100,",
    "                      'personalizacao':'bordado', 'tecido':'algodao_basico'},",
    "        },",
    "        ...",
    "    ],",
    "    'ultimo_assunto': 'combinado_preco_qtd_produto',",
    "    'ativa': True,",
    "}",
])

h(doc, "Regras de invalidação no foco_atual (não no histórico)", level=3)

t2 = doc.add_table(rows=4, cols=2)
t2.style = "Light Grid Accent 1"
t2.rows[0].cells[0].text = ""
add_run(t2.rows[0].cells[0].paragraphs[0], "Slot que mudou", bold=True)
t2.rows[0].cells[1].text = ""
add_run(t2.rows[0].cells[1].paragraphs[0], "O que zera no foco", bold=True)
dados_t2 = [
    ("produto", "personalização, cor, grade (propriedades do produto anterior)"),
    ("numero_pedido", "quantidade (provável ano do pedido bagunçando)"),
    ("Intenção mudou de orçamento pra status", "produto e quantidade do orçamento"),
]
for i, (k, v) in enumerate(dados_t2, 1):
    t2.rows[i].cells[0].text = k
    t2.rows[i].cells[1].text = v

p(doc)
h(doc, "Três opções de implementação", level=3)
t3 = doc.add_table(rows=4, cols=4)
t3.style = "Light Grid Accent 1"
hd = ["Opção", "Cobertura", "Esforço", "Risco"]
for j, lbl in enumerate(hd):
    t3.rows[0].cells[j].text = ""
    add_run(t3.rows[0].cells[j].paragraphs[0], lbl, bold=True)
opcoes = [
    ("(a) Só foco com invalidação", "Não mistura. Não revisita.", "Baixo", "Baixo"),
    ("(b) Foco + histórico passivo ⭐", "Não mistura. Memória existe pra demo. Bot não navega ativamente.", "Médio", "Baixo"),
    ("(c) Foco + histórico ativo + comandos", "Bot entende 'volta pros polos'.", "Alto", "Médio (pode regredir cenários ok hoje)"),
]
for i, row in enumerate(opcoes, 1):
    for j, val in enumerate(row):
        t3.rows[i].cells[j].text = val

p(doc)
rec = p(doc)
add_run(rec, "Recomendação: ", bold=True)
rec.add_run("opção (b). Resolve a coerência (foco bem gerenciado), guarda memória completa, e "
            "dá uma 'feature de apresentação' linda — mostrar o histórico pra professora como "
            "prova viva de que o bot lembra de tudo. (c) fica de sprint futura.")

# ===== SECAO 4: plano priorizado =====
doc.add_page_break()
h(doc, "5. Plano de correção priorizado", level=1)

h(doc, "Bloco A — Indispensável", level=2, color=COR_CRITICO)
p(doc, "Resolve 7 dos 8 críticos + 2 altos. Esforço estimado: 3-4 horas.")
itens_a = [
    "Reescrever bot/contexto.py com 'foco_atual + historico_turnos' e regras de invalidação (CRÍTICO 1, 22, ALTO 23)",
    "Fix do regex de quantidade pra não pegar ano de pedido (CRÍTICO 2)",
    "Não extrair quantidade quando há aguardando_opcao (CRÍTICO 3)",
    "Reescrever is_despedida com match restritivo (CRÍTICO 4)",
    "numero_pedido é slot de turno único (CRÍTICO 5)",
    "Regras do classifier por verbo antes de por slot: 'cancelar' antes de status (CRÍTICO 6)",
    "Detectar negação ('sem bordado') no extractor (CRÍTICO 22)",
    "Sessão por usuário no Gradio via gr.State() (CRÍTICO 7)",
]
for it in itens_a:
    doc.add_paragraph(it, style="List Number")

h(doc, "Bloco B — Recomendado se sobrar tempo", level=2, color=COR_ALTO)
itens_b = [
    "Refatorar extrair_slots pra função pura (DESIGN 21)",
    "Follow-up sem número herda intenção (ALTO 9)",
    "selecao_opcao não vira ultimo_assunto (ALTO 10)",
    "qtd_grande_volume não come viabilidade_producao (ALTO 11)",
    "Produto inexistente vira pergunta clarificadora (ALTO 30)",
    "Múltiplas personalizações: bot pergunta qual (ALTO 25)",
]
for it in itens_b:
    doc.add_paragraph(it, style="List Number")

h(doc, "Bloco C — Polimento", level=2, color=COR_MEDIO)
itens_c = [
    "Plurais soltos (MÉDIO 13)",
    "Conflito quantidade/prazo (ALTO 8)",
    "Letra M ambígua (MÉDIO 16)",
    "Slot urgente sem uso (MÉDIO 17)",
    "Fuzzy match em produto — 'camizetas' sobrevive (MÉDIO 27)",
    "Singular/plural correto (MÉDIO 26)",
    "Número de pedido formato alternativo (MÉDIO 28)",
    "status_pedido usar lookup_pedidos.csv (MÉDIO 29)",
]
for it in itens_c:
    doc.add_paragraph(it, style="List Number")

h(doc, "Bloco D — Não mexer agora", level=2, color=COR_BAIXO)
itens_d = [
    "Substring em silk (raro)",
    "Slot dinâmico do menu (não atrapalha)",
    "Datas obsoletas no CSV (só editar antes de apresentar)",
    "Inconsistência slots.csv (documentação)",
    "Dois números de pedido na mesma frase (BAIXO 32)",
    "Quantidade extrapolada 99999 (MÉDIO 31)",
    "Bug 'do bem' 33 (informativo, não acionável)",
]
for it in itens_d:
    doc.add_paragraph(it, style="List Bullet")

# ===== SECAO 5: gradio melhorias =====
h(doc, "6. Sugestões pro front_gradio.py (visual)", level=1)
sug = [
    "Avatar do usuário (hoje só o bot tem)",
    "Cabeçalho com logo e nome da empresa",
    "Exemplos focados em produção: 'preço de 100 polos com bordado', 'consigo produzir 500 peças em 10 dias?'",
    "Botão 'limpar conversa' (importante depois de resolver CRÍTICO 7)",
    "Painel lateral mostrando 'memória ativa' (slots + histórico) — feature linda pra apresentação",
    "Trocar paleta pink/rose por algo que case melhor com confecção (terra, verde-oliva) — opcional",
]
for s in sug:
    doc.add_paragraph(s, style="List Bullet")

# ===== SECAO 6: verificação =====
h(doc, "7. Bateria de verificação (15 cenários)", level=1)
p(doc, "Após implementar Bloco A, todos devem passar — sem resposta absurda e sem perder contexto.")

testes = [
    "['preço de 200 polos', 'e com bordado?', 'e o prazo?']",
    "['preço de 100 polos com bordado', 'agora 200 calças jeans', 'quanto custa?']",
    "['preço de 100 polos', 'qualidade', '1', 'e o prazo?']",
    "['preço de 100 polos', 'obrigado, agora me fala o prazo']",
    "['meu pedido FF-2024-0123 está em qual etapa']",
    "['quero cancelar meu pedido FF-2025-0001']",
    "['pedido FF-2025-0001', 'tem algodão em estoque?']",
    "['preço de 200 polos', 'e em algodão?']",
    "['consigo 1000 peças de moletom em 10 dias?']",
    "['preço de 30 vestidos']",
    "['preço de 30 calças jeans']",
    "['preciso em 100 dias']",
    "['vc é burro', 'desculpa, queria saber preço de 50 camisetas']",
    "['preço de 100 polos', 'ok', 'e o prazo?']",
    "['urgente! preciso 50 camisetas em 7 dias']",
]
code_block(doc, [f"  {i+1:>2}. {t}" for i, t in enumerate(testes)])

# ===== SECAO 7: glossário =====
doc.add_page_break()
h(doc, "8. Glossário (pra você aprender os termos)", level=1)

t4 = doc.add_table(rows=13, cols=3)
t4.style = "Light Grid Accent 1"
for j, lbl in enumerate(["Termo", "O que é", "Exemplo"]):
    t4.rows[0].cells[j].text = ""
    add_run(t4.rows[0].cells[j].paragraphs[0], lbl, bold=True)
glos = [
    ("Intenção", "O 'objetivo' do que o usuário quer", "'Quanto custa?' → combinado_preco_qtd_produto"),
    ("Slot", "Uma 'variável' extraída da mensagem", "'200 polos' → quantidade=200, produto=polo"),
    ("Extração", "Achar slots dentro do texto", "Função extrair_slots"),
    ("Classificação", "Decidir qual intenção é", "Função classificar"),
    ("Contexto / sessão", "O que o bot lembra ao longo da conversa", "Dict sessao no código"),
    ("Merge / mesclagem", "Juntar slots novos com os antigos", "Função merge_slots"),
    ("Regex", "Padrão pra achar texto", r"\b(\d+)\b = 'um número solto'"),
    ("Word boundary (\\b)", "'Borda de palavra' no regex", r"\bla\b casa 'la' mas não casa 'lavanderia'"),
    ("Substring", "Pedaço dentro de outra string", "'silk' é substring de 'silkada'"),
    ("Fuzzy match", "Comparação por similaridade", "'moletoom' vs 'moletom' tem score ~95"),
    ("Fallback", "Resposta padrão quando o bot não entende", "intencao=fallback"),
    ("Side effect", "Função que muda algo além de retornar", "extrair_slots mexe na sessao"),
]
for i, (t, d, e) in enumerate(glos, 1):
    t4.rows[i].cells[0].text = t
    t4.rows[i].cells[1].text = d
    t4.rows[i].cells[2].text = e

# ===== SECAO 8: próximos passos =====
h(doc, "9. Próximos passos", level=1)
passos = [
    "Você lê este documento e revisa os bugs/severidades.",
    "Testamos a interface real (Gradio) pra confirmar os críticos com cliques.",
    "Implementamos o Bloco A inteiro (3-4 horas estimadas).",
    "Rodamos os 15 cenários de verificação.",
    "Decidimos se vamos pro Bloco B/C ou pra melhorar o front.",
    "Commitamos as mudanças com mensagens claras pra mostrar pra equipe e professora.",
]
for s in passos:
    doc.add_paragraph(s, style="List Number")

p(doc)
fim = p(doc)
fim.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(fim, "— fim do documento —", italic=True, color=COR_CINZA)

doc.save(OUT)
print(f"OK: {OUT}")
print(f"Tamanho: {OUT.stat().st_size} bytes")
