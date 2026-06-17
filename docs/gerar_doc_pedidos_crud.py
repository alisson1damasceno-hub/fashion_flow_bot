"""
Gera proposta_crud_pedidos.docx — proposta para a equipe sobre como
implementar o CRUD de pedidos + a nova tabela pedidos.csv, mantendo o
bot fiel ao setor de produção.

Execute: python docs/gerar_doc_pedidos_crud.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

OUT = Path(__file__).parent / "proposta_crud_pedidos.docx"

ROXO = RGBColor(0x76, 0x4A, 0xB0)
VERDE = RGBColor(0x27, 0xAE, 0x60)
LARANJA = RGBColor(0xE6, 0x7E, 0x22)
VERMELHO = RGBColor(0xC0, 0x39, 0x2B)
CINZA = RGBColor(0x55, 0x55, 0x55)


def add_run(p, text, bold=False, italic=False, color=None, mono=False, size=None):
    r = p.add_run(text)
    if bold: r.bold = True
    if italic: r.italic = True
    if color: r.font.color.rgb = color
    r.font.name = "Consolas" if mono else "Calibri"
    if mono and size is None: size = 9.5
    if size: r.font.size = Pt(size)
    return r


def h(doc, text, level=1, color=None):
    p = doc.add_heading(level=level)
    r = p.add_run(text)
    if color: r.font.color.rgb = color
    return p


def p(doc, text=""):
    return doc.add_paragraph(text)


def code(doc, lines):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(0.5)
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    add_run(para, lines if isinstance(lines, str) else "\n".join(lines),
            mono=True, color=CINZA, size=9.5)
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'F4F4F4')
    pPr.append(shd)


def tabela(doc, headers, linhas):
    """Cria tabela com cabeçalho em negrito + linhas. Retorna a tabela."""
    tab = doc.add_table(rows=len(linhas) + 1, cols=len(headers))
    tab.style = "Light Grid Accent 1"
    for j, lbl in enumerate(headers):
        tab.rows[0].cells[j].text = ""
        add_run(tab.rows[0].cells[j].paragraphs[0], lbl, bold=True)
    for i, row in enumerate(linhas, 1):
        for j, val in enumerate(row):
            tab.rows[i].cells[j].text = str(val)
    return tab


# ============================================================
doc = Document()
for s in doc.sections:
    s.left_margin = Cm(2.0); s.right_margin = Cm(2.0)
    s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)

# CAPA
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(t, "\n\n\nFashion Flow Bot\n", bold=True, size=28, color=ROXO)
add_run(t, "Proposta: CRUD de Pedidos + tabela pedidos.csv\n", bold=True, size=18)

s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(s, "\nComo registrar, consultar, alterar e cancelar pedidos\n"
           "sem o bot sair do papel de setor de produção\n",
        italic=True, color=CINZA)
add_run(s, "Bia Costa • 2026-06-16\n", color=CINZA, size=10)

p(doc)
caixa = doc.add_paragraph()
caixa.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(caixa, "Documento de discussão — precisa de 1 decisão da equipe antes de codar (seção 10)",
        bold=True, size=11)

doc.add_page_break()

# 1. Objetivo
h(doc, "1. O que precisamos construir", level=1)
p(doc,
  "A próxima entrega é dar ao bot a capacidade de lidar com pedidos de verdade — "
  "um CRUD (Create, Read, Update, Delete) — em cima dos CSVs que já temos. O pedido da entrega:")
metas = [
    "CRUD de pedidos coerente com as tabelas CSV atuais.",
    "Criar a nova tabela pedidos.csv.",
    "Definir bem as etapas, batendo com as que já existem.",
    "Registrar um pedido na tabela nova, gerando um ID automaticamente.",
    "Reconhecer um pedido específico pelo ID que o cliente mandar.",
    "Ter uma intenção dedicada para perguntar ao cliente o ID do pedido.",
    "Adicionar o status do pedido.",
]
for m in metas:
    doc.add_paragraph(m, style="List Bullet")
restr = p(doc)
add_run(restr, "Restrição que guia tudo: ", bold=True)
restr.add_run("somos APENAS o bot do setor de produção. Isso muda o que o bot pode "
              "responder (e o que ele encaminha pra outro setor).")

# 2. A descoberta que muda o desenho
h(doc, "2. A descoberta que muda o desenho", level=1)
p(doc,
  "Antes de propor qualquer coisa, abri os arquivos. O ponto mais importante:")
achado = p(doc)
add_run(achado, "O arquivo lookup_pedidos.csv NÃO é uma tabela de pedidos. ", bold=True, color=VERMELHO)
achado.add_run("Ele é uma tabela de ETAPAS. As colunas são "
               "etapa, pode_alterar, descricao, observacao — e as 6 linhas são as fases "
               "de produção (modelagem, corte, costura, personalizacao, qualidade, "
               "embalagem_expedicao). Não há nenhum pedido real, nenhum ID, nenhum cliente.")
p(doc, "Consequências disso no código de hoje:")
doc.add_paragraph(
    "As intenções status_pedido, alterar_pedido_especifico e cancelar_pedido "
    "(responder.py) nunca reconhecem um pedido real — elas só leem a lista de etapas "
    "e respondem genérico ('fale com a logística / vendas'). O slot numero_pedido "
    "(formato FF-AAAA-NNNN) é extraído, mas não existe onde procurar esse ID.",
    style="List Bullet")
doc.add_paragraph(
    "Já existe uma regra de fronteira de setor embutida: a intenção etapa_consulta diz "
    "literalmente que 'o setor de produção não fornece status em tempo real — isso é da "
    "logística'. Ou seja: hoje o bot SE RECUSA a dar status. Nossa meta de 'adicionar "
    "status' parece contradizer isso — e é justamente essa contradição que a próxima "
    "seção resolve.",
    style="List Bullet")

# 3. Ideia central
h(doc, "3. Ideia central: status = etapa de produção", level=1)
p(doc,
  "A forma de adicionar status SEM invadir a logística é separar dois tipos de status:")
tabela(doc,
       ["Tipo de status", "De quem é", "O bot responde?"],
       [["Status de fabricação (em que etapa da fábrica a peça está)", "Produção (nós)", "SIM"],
        ["Status de entrega (saiu pra entrega? cadê meu pacote? previsão de chegada)", "Logística", "Não — encaminha"]])
p(doc)
rec = p(doc)
add_run(rec, "Decisão de design: ", bold=True)
rec.add_run("o 'status do pedido' que o bot dá = a etapa de produção atual. "
            "Isso é exatamente o que a fábrica sabe e não pisa no território da logística. "
            "Resolve a contradição da seção 2.")

# 4. Arquitetura
h(doc, "4. Arquitetura: duas tabelas em vez de uma", level=1)
p(doc,
  "Seguindo o padrão que já usamos (tabela de referência + tabela de dado), separo "
  "a estrutura em duas:")
tabela(doc,
       ["Tabela", "Papel", "Conteúdo"],
       [["lookup_etapas.csv\n(renomear o atual lookup_pedidos.csv)", "referência / dimensão",
         "As 6 etapas + pode_alterar + descrição. Quase nunca muda."],
        ["pedidos.csv\n(NOVO)", "dado / fato",
         "Os pedidos reais, cada um com ID e apontando para a etapa atual."]])
p(doc, "Como as duas tabelas e o pipeline se conectam:")
code(doc, [
    "intencoes.csv ──► classifier ──► responder",
    "                                     │",
    "          ┌──────────────────────────┴───────────────────────┐",
    "          ▼                                                   ▼",
    "  lookup_etapas.csv   ◄──  FK (etapa_atual)  ──         pedidos.csv",
    "  (6 etapas, regra                                 (pedidos reais: ID,",
    "   pode_alterar)                                     etapa, status, datas)",
])
p(doc, "O pedido aponta para a etapa pelo campo etapa_atual; assim a regra "
       "'pode alterar?' fica num lugar só (a tabela de etapas).")

# 5. Schema do pedidos.csv
doc.add_page_break()
h(doc, "5. Schema proposto do pedidos.csv", level=1)
p(doc, "Reaproveitando o vocabulário de slots que o extractor já produz "
       "(camiseta_basica, algodao_pima, preto...), pra ficar consistente com o resto:")
tabela(doc,
       ["Coluna", "Tipo", "Exemplo", "Origem / observação"],
       [["numero_pedido", "ID (PK)", "FF-2026-0001", "formato do slot numero_pedido"],
        ["data_criacao", "data", "2026-06-16", "gerada no Create"],
        ["produto", "enum", "camiseta_basica", "slot produto"],
        ["quantidade", "inteiro", "150", "slot quantidade"],
        ["cor", "enum", "preto", "slot cor"],
        ["tamanho", "enum", "M", "slot tamanho / grade"],
        ["tecido", "enum", "algodao_pima", "slot tecido"],
        ["personalizacao", "enum", "bordado / nenhuma", "slot personalizacao"],
        ["etapa_atual", "enum (FK → lookup_etapas)", "corte", "É O STATUS de fabricação"],
        ["status", "enum", "em_producao / concluido / cancelado / pausado", "ciclo de vida macro"],
        ["data_prevista", "data", "2026-07-10", "previsão de término da produção"],
        ["observacao", "texto", "(livre)", "anotações internas"]])
p(doc)
nota = p(doc)
add_run(nota, "Dois níveis de status. ", bold=True)
nota.add_run("etapa_atual é o granular (o que o cliente quer saber). status é o macro "
             "(serve pro Update/Delete terem o que mexer) — e status='cancelado' vira o "
             "'delete suave' que preserva o histórico.")
p(doc, "Exemplo de uma linha do arquivo:")
code(doc, [
    "numero_pedido,data_criacao,produto,quantidade,cor,tamanho,tecido,",
    "personalizacao,etapa_atual,status,data_prevista,observacao",
    "FF-2026-0001,2026-06-16,camiseta_basica,150,preto,M,algodao_pima,",
    "bordado,corte,em_producao,2026-07-10,cliente PJ",
])

# 6. As 6 etapas
h(doc, "6. As etapas (que já estão definidas)", level=1)
p(doc, "Mantemos exatamente as 6 etapas que o CSV atual já tem — só mudamos o nome do "
       "arquivo para lookup_etapas.csv. A coluna pode_alterar é a regra de negócio:")
tabela(doc,
       ["Etapa", "Pode alterar?", "O que está acontecendo"],
       [["modelagem", "SIM", "Ficha técnica e moldes. Ainda dá pra mudar produto, cor, tamanho e quantidade."],
        ["corte", "não", "Tecido sendo cortado conforme os moldes. Alterações já inviáveis."],
        ["costura", "não", "Peças cortadas sendo montadas pelas costureiras."],
        ["personalizacao", "não", "Bordado, silk ou DTF sendo aplicados."],
        ["qualidade", "não", "Inspeção final dimensional e visual."],
        ["embalagem_expedicao", "não", "Peças aprovadas, embaladas e prontas para envio."]])
p(doc)
fluxo = p(doc)
add_run(fluxo, "Fluxo (este também é o avanço de status no Update): ", bold=True)
add_run(fluxo, "modelagem → corte → costura → personalizacao → qualidade → embalagem_expedicao",
        mono=True, size=9.5)

# 7. CRUD mapeado ao setor
doc.add_page_break()
h(doc, "7. O CRUD, mapeado ao setor de produção", level=1)
tabela(doc,
       ["Operação", "O que faz", "Quem usa", "Onde liga no código atual"],
       [["C — Registrar", "Coleta os slots (produto, qtd, cor...), GERA O ID e grava a linha "
         "nova com etapa_atual=modelagem e status=em_producao.", "operador da produção",
         "nova intenção registrar_pedido + preenchimento dos slots"],
        ["R — Consultar", "Cliente manda o ID → o bot acha a linha → responde "
         "produto / quantidade / etapa atual / previsão / se ainda dá pra alterar.",
         "cliente (principal)", "reescrever status_pedido pra ler o pedidos.csv de verdade"],
        ["U — Atualizar", "Avançar a etapa (modelagem→corte→...), alterar campos (só se "
         "pode_alterar=sim) ou pausar.", "operador / cliente",
         "alterar_pedido_especifico passa a validar contra o pedido real"],
        ["D — Cancelar", "status=cancelado (soft delete: preserva o histórico em vez de "
         "apagar a linha).", "operador / cliente", "cancelar_pedido grava de verdade"]])
p(doc)
sd = p(doc)
add_run(sd, "Por que soft delete: ", bold=True)
sd.add_run("marcar 'cancelado' em vez de apagar a linha mantém o histórico — mais realista "
           "pra uma fábrica e mais seguro pra demo (nada some).")

# 8. ID + intenção de perguntar o ID
h(doc, "8. Geração do ID e a intenção de perguntar o ID", level=1)
p(doc, "O ID segue o formato que o slot numero_pedido já reconhece: FF-AAAA-NNNN, "
       "sequencial por ano. Geramos pegando o maior NNNN do ano e somando 1 "
       "(hoje o primeiro seria FF-2026-0001):")
code(doc, [
    "def gerar_id(df_pedidos, ano):",
    "    prefixo = f'FF-{ano}-'",
    "    do_ano = df_pedidos[df_pedidos['numero_pedido'].str.startswith(prefixo)]",
    "    if do_ano.empty:",
    "        seq = 1",
    "    else:",
    "        seq = do_ano['numero_pedido'].str[-4:].astype(int).max() + 1",
    "    return f'{prefixo}{seq:04d}'   # → FF-2026-0001",
])
p(doc)
inten = p(doc)
add_run(inten, "A intenção de perguntar o ID ", bold=True)
inten.add_run("já existe pela metade: o follow-up de status_pedido é 'Pode me informar o "
              "número do pedido? Ex: FF-2025-0389'. Falta um estado de espera, igual ao "
              "aguardando_opcao dos menus. O fluxo fica:")
doc.add_paragraph("Cliente pergunta do pedido SEM ID → bot pergunta o ID e marca "
                  "aguardando_numero_pedido = True.", style="List Number")
doc.add_paragraph("Cliente responde FF-2026-0001 → o extractor já captura o ID → bot consulta "
                  "a linha.", style="List Number")
doc.add_paragraph("Validação: achou → mostra o status; formato errado → reexplica o formato; "
                  "formato ok mas inexistente → 'não encontrei o FF-2026-9999, confere o "
                  "número?'.", style="List Number")

# 9. O que muda no código
h(doc, "9. O que muda no código", level=1)
tabela(doc,
       ["Arquivo", "Mudança"],
       [["bot/loader.py", "adicionar 'etapas': 'lookup_etapas.csv' e 'pedidos': 'pedidos.csv'"],
        ["bot/responder.py", "trocar os dados['pedidos'] que hoje leem etapas por "
         "dados['etapas']; e implementar a leitura real do pedido por ID"],
        ["bot/pedidos_repo.py (NOVO)", "módulo só pra ler/gravar o pedidos.csv: gerar_id, "
         "criar, atualizar, cancelar. Separa a persistência das lookups."],
        ["bot/contexto.py", "novo estado aguardando_numero_pedido (e, se for ter Create "
         "conversacional, um estado de preenchimento do pedido)"],
        ["data/intencoes.csv", "nova intenção registrar_pedido; ajustar os textos de "
         "status_pedido / etapa_consulta pra refletir que produção agora dá a etapa de "
         "fabricação (mas entrega continua na logística)"],
        ["data/lookup_etapas.csv", "renomear o atual lookup_pedidos.csv (conteúdo igual)"],
        ["data/pedidos.csv (NOVO)", "a tabela de pedidos reais, com alguns registros de "
         "exemplo (seed) pra testar a consulta"]])

# 10. Decisão da equipe
doc.add_page_break()
h(doc, "10. Decisão que precisamos bater o martelo", level=1)
p(doc, "O desenho acima é consenso. O que muda o tamanho do trabalho é UMA pergunta: "
       "como o CRUD aparece na conversa — todo mundo faz tudo, ou separamos papéis?")
tabela(doc,
       ["Opção", "Como funciona", "Trade-off"],
       [["A) Operador + cliente (recomendada)", "Cliente só CONSULTA por ID (R). Registrar, "
         "avançar e cancelar (C/U/D) são comandos de operador (ex: /novo, /avancar, /cancelar).",
         "Mais realista pro setor; mostra os 4 do CRUD com controle de acesso."],
        ["B) Tudo por linguagem natural", "Sem separar papéis: qualquer um registra, consulta, "
         "altera e cancela conversando.", "Mais simples de codar; menos realista."],
        ["C) Só cliente-facing", "Registrar + consultar + cancelar por conversa; sem painel de "
         "operador.", "Update fica limitado a cancelar."]])
p(doc)
p2 = p(doc)
add_run(p2, "Minha recomendação: ", bold=True)
p2.add_run("opção A. É a mais fiel ao 'somos só produção' e ainda demonstra o CRUD inteiro "
           "de um jeito organizado. Mas é decisão da equipe.")
p(doc)
p3 = p(doc)
add_run(p3, "Pergunta menor (já com sugestão): ", bold=True)
p3.add_run("o Create faz parte da produção? Sim — a produção registra a ordem de fabricação "
           "quando ela entra na fábrica. Se a equipe achar que o pedido 'vem pronto' de "
           "vendas, a gente só popula a tabela com um seed e o bot foca em R/U/D.")

# 11. Próximos passos
h(doc, "11. Próximos passos", level=1)
passos = [
    "A equipe lê este documento e escolhe a opção da seção 10 (A, B ou C).",
    "Renomear lookup_pedidos.csv → lookup_etapas.csv e ajustar o loader/responder.",
    "Criar pedidos.csv (com 3-5 pedidos de exemplo) e o módulo pedidos_repo.py.",
    "Fazer o R (consultar por ID) funcionar de ponta a ponta — é o núcleo da entrega.",
    "Encaixar C/U/D conforme a opção escolhida.",
    "Testar: consultar um ID que existe, um que não existe e um com formato errado.",
]
for s in passos:
    doc.add_paragraph(s, style="List Number")

p(doc)
fim = p(doc)
fim.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(fim, "— fim do documento —", italic=True, color=CINZA)

doc.save(OUT)
print(f"OK: {OUT}")
print(f"Tamanho: {OUT.stat().st_size} bytes")
