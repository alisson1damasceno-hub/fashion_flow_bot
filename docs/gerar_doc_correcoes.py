"""
Gera correcoes_aplicadas.docx — explica TUDO que foi mudado no Bloco A,
no estilo "antes / depois / por que", didático pra estudante de 1º sem.

Execute: python docs/gerar_doc_correcoes.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

OUT = Path(__file__).parent / "correcoes_aplicadas.docx"

ROXO   = RGBColor(0x76, 0x4A, 0xB0)
ANTES  = RGBColor(0xC0, 0x39, 0x2B)  # vermelho
DEPOIS = RGBColor(0x27, 0xAE, 0x60)  # verde
CINZA  = RGBColor(0x55, 0x55, 0x55)
AZUL   = RGBColor(0x2C, 0x6F, 0xB5)


def add_run(p, text, bold=False, italic=False, color=None, mono=False, size=None):
    r = p.add_run(text)
    if bold: r.bold = True
    if italic: r.italic = True
    if color: r.font.color.rgb = color
    r.font.name = "Consolas" if mono else "Calibri"
    if mono and size is None: size = 9.5
    if size: r.font.size = Pt(size)
    return r


def h(doc, text, level=1, color=None):
    p = doc.add_heading(level=level)
    r = p.add_run(text)
    if color: r.font.color.rgb = color
    return p


def p(doc, text=""):
    return doc.add_paragraph(text)


def code_block(doc, lines, cor_fundo="F4F4F4"):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(0.5)
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    add_run(para, lines if isinstance(lines, str) else "\n".join(lines),
            mono=True, color=CINZA, size=9.5)
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), cor_fundo)
    pPr.append(shd)
    return para


def antes_depois(doc, antes_lines, depois_lines):
    """Bloco visual com 'antes' (vermelho) e 'depois' (verde)."""
    px = p(doc)
    add_run(px, "🔴 Antes:", bold=True, color=ANTES)
    code_block(doc, antes_lines, cor_fundo="FFEEEC")
    px = p(doc)
    add_run(px, "🟢 Depois:", bold=True, color=DEPOIS)
    code_block(doc, depois_lines, cor_fundo="EEFBEF")


# ============================================================
# DOCUMENTO
# ============================================================

doc = Document()
for s in doc.sections:
    s.left_margin = Cm(2.0); s.right_margin = Cm(2.0)
    s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)

# ---------- CAPA ----------
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(t, "\n\n\nFashion Flow Bot\n", bold=True, size=28, color=ROXO)
add_run(t, "Correções Aplicadas (Bloco A)\n", bold=True, size=18)

s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(s, "\nExplicação didática de tudo que mudou no código\n", italic=True, color=CINZA)
add_run(s, "Bia Costa • 2026-06-10\n", color=CINZA, size=10)

p(doc)
caixa = p(doc)
caixa.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(caixa, "8 críticos resolvidos  •  arquitetura nova  •  15/15 cenários passando",
        bold=True, size=11)
doc.add_page_break()

# ---------- INTRO ----------
h(doc, "1. O que foi feito", level=1)
p(doc,
  "Esse documento é a versão narrada do Bloco A da auditoria. Pra cada mudança eu mostro: "
  "o que tava errado (antes), como ficou (depois), e por que essa decisão. "
  "Tudo em PT-BR e em ordem de pra você poder ler enquanto compara com o código.")

p(doc,
  "Resumo: o problema central era que o bot misturava informações de diferentes momentos da "
  "conversa. Eu reescrevi o módulo de contexto com um design novo: separei 'foco atual' "
  "(o que está sendo discutido agora) de 'histórico de turnos' (memória completa). "
  "Daí em diante, tudo passou a fazer sentido.")

# ---------- VISÃO GERAL ----------
h(doc, "2. Visão geral da nova arquitetura", level=1)

p(doc, "Pense numa folha de pedido sendo escrita à caneta:")
sub = p(doc)
add_run(sub, "• Foco atual ", bold=True)
sub.add_run("é a linha que o atendente está preenchendo agora. Se você mudar de produto, "
            "ele risca o que não vale mais e escreve o novo.")
sub = p(doc)
add_run(sub, "• Histórico de turnos ", bold=True)
sub.add_run("é a folha cheia, sem rasura. Tudo o que foi conversado fica registrado, "
            "mesmo que não esteja mais 'em pauta'.")

code_block(doc, [
    "sessao = {",
    "    'foco_atual': {",
    "        'produto': 'polo',",
    "        'quantidade': 200,",
    "        'personalizacao': 'bordado',",
    "    },",
    "    'historico_turnos': [",
    "        {'msg': 'preço de 200 polos',",
    "         'intencao': 'combinado_preco_qtd_produto',",
    "         'slots': {...}, 'resposta': '...'},",
    "        ...",
    "    ],",
    "    'ultimo_assunto': 'combinado_preco_qtd_produto',",
    "    'aguardando_opcao': None,",
    "    'ativa': True,",
    "}",
])

p(doc, "Quando o usuário muda de produto, o foco se reorganiza (esquece personalização/cor/tecido), "
       "mas o histórico continua intacto. Se um dia a gente implementar 'voltar pro orçamento "
       "anterior', a memória já está lá.")

# ---------- BUGS ----------
doc.add_page_break()
h(doc, "3. Cada bug consertado, com antes e depois", level=1)

# ===== BUG 1 =====
h(doc, "3.1  CRÍTICO 1 — Slots não grudam mais quando muda produto", level=2, color=ROXO)
p(doc, "Era o bug mais fundamental. O bot misturava informações de diferentes orçamentos.")

antes_depois(doc, [
    "Você: preço de 100 polos com bordado em algodão",
    "Bot:  Para 100 polos com bordado: R$ 88/peça",
    "Você: agora 200 calças jeans",
    "Bot:  Para 200 PEÇAS DE POLO COM BORDADO: R$ 84/peça  ← errado!",
], [
    "Você: preço de 100 polos com bordado em algodão",
    "Bot:  Para 100 polos com bordado: R$ 88/peça",
    "Você: agora 200 calças jeans",
    "Bot:  Para 200 peças de calça jeans: R$ 125/peça  ✓",
])

px = p(doc)
add_run(px, "O que mudou no código: ", bold=True)
px.add_run("bot/contexto.py tem agora a função ")
add_run(px, "merge_com_contexto", mono=True)
px.add_run(", com uma regra de invalidação:")

code_block(doc, [
    "# se o produto mudou no turno atual,",
    "# esquece slots-filhos do produto antigo",
    "produto_novo = slots_do_turno.get('produto')",
    "if produto_novo and foco.get('produto') and produto_novo != foco['produto']:",
    "    for k in SLOTS_FILHOS_PRODUTO:   # personalizacao, cor, grade, tecido",
    "        foco.pop(k, None)",
])

px = p(doc)
add_run(px, "Por que essa regra: ", bold=True)
px.add_run("personalização, cor, grade e tecido são propriedades específicas de UM produto. "
           "Bordado num polo é uma decisão; bordado numa calça jeans é outra. Quando o cliente "
           "muda de produto, é razoável a gente perguntar de novo (ou começar limpo).")

# ===== BUG 2 =====
h(doc, "3.2  CRÍTICO 2 — Quantidade não pega mais o ano do pedido", level=2, color=ROXO)

antes_depois(doc, [
    "Você: meu pedido FF-2024-0123 está em qual etapa",
    "Bot:  [slots] quantidade=2024, numero_pedido=FF-2024-0123",
    "      (e o 2024 envenenava os próximos turnos)",
], [
    "Você: meu pedido FF-2024-0123 está em qual etapa",
    "Bot:  [slots] numero_pedido=FF-2024-0123  (sem quantidade)",
])

px = p(doc)
add_run(px, "O que mudou: ", bold=True)
px.add_run("o extractor agora extrai o número de pedido PRIMEIRO e remove esse trecho da string "
           "antes de procurar quantidade. Além disso, a quantidade agora exige uma unidade "
           "obrigatória (peças, unidades, polos, camisetas etc):")

code_block(doc, [
    "# 1) extrai numero_pedido e remove da string",
    "match_pedido = re.search(r'FF-\\d{4}-\\d{4}', mensagem, re.IGNORECASE)",
    "if match_pedido:",
    "    slots['numero_pedido'] = match_pedido.group(0).upper()",
    "    t = re.sub(r'ff-\\d{4}-\\d{4}', ' ', t)",
    "",
    "# 2) quantidade exige unidade — não pega número solto",
    "match_qtd = re.search(rf'\\b(\\d+)\\s*({UNIDADES_QTD})\\b', t)",
])

# ===== BUG 3 =====
h(doc, "3.3  CRÍTICO 3 — Menu não corrompe mais a quantidade", level=2, color=ROXO)

antes_depois(doc, [
    "Você: preço de 100 polos               → quantidade=100",
    "Você: qualidade                         → menu mostrado",
    "Você: 1                                  → escolheu opção 1",
    "Você: e o prazo?",
    "Bot:  Para 1 PEÇA de polo: 18-22 dias   ← '1' virou quantidade!",
], [
    "Você: preço de 100 polos               → quantidade=100",
    "Você: qualidade                         → menu mostrado",
    "Você: 1                                  → escolheu opção 1",
    "Você: e o prazo?",
    "Bot:  Para 100 peças de polo: 22-26 dias  ✓",
])

px = p(doc)
add_run(px, "O que mudou: ", bold=True)
px.add_run("a função extrair_slots ganhou um parâmetro 'em_menu'. Quando o bot está aguardando "
           "uma escolha de menu, o extractor não tenta extrair quantidade — assim o '1' não vira slot.")

code_block(doc, [
    "def extrair_slots(mensagem, em_menu=False):",
    "    if em_menu:",
    "        return {}   # número digitado no menu não é slot",
    "    ...",
])

# ===== BUG 4 =====
h(doc, "3.4  CRÍTICO 4 — 'obrigado, agora me fala...' não reseta mais a sessão", level=2, color=ROXO)

antes_depois(doc, [
    "Você: preço de 100 polos",
    "Você: obrigado, agora me fala o prazo",
    "Bot:  Até logo!  ← a sessão foi RESETADA antes de processar",
], [
    "Você: preço de 100 polos",
    "Você: obrigado, agora me fala o prazo",
    "Bot:  Para 100 peças de polo: 22-26 dias úteis  ✓",
])

px = p(doc)
add_run(px, "O que mudou: ", bold=True)
px.add_run("a função is_despedida ficou MUITO mais restritiva. Antes ela usava substring match "
           "(qualquer 'obrigado' no meio do texto contava); agora ela só dispara se a mensagem "
           "INTEIRA for uma despedida curta:")

code_block(doc, [
    "def is_despedida(mensagem):",
    "    t = normalizar(mensagem).strip()",
    "    if t in _DESPEDIDAS_EXATAS:  # 'tchau', 'até logo', 'flw'...",
    "        return True",
    "    # ou: mensagem ≤ 3 palavras, todas casuais ou de despedida",
    "    palavras = re.findall(r'\\w+', t)",
    "    if 1 <= len(palavras) <= 3 and all(...):",
    "        return True",
    "    return False",
])

px = p(doc)
add_run(px, "Além disso, agradecimentos ('obrigado', 'valeu') saíram da lista de despedida e "
            "foram pra is_casual — eles continuam a conversa, não encerram.", italic=True)

# ===== BUG 5 e 6 =====
h(doc, "3.5  CRÍTICO 5 e 6 — Número de pedido não trava em status, e 'cancelar' funciona", level=2, color=ROXO)

p(doc, "Os dois bugs vinham da mesma raiz: a regra do classifier que checava o slot "
       "numero_pedido era avaliada antes das regras por verbo.")

antes_depois(doc, [
    "Você: pedido FF-2025-0001",
    "Você: tem algodão em estoque?",
    "Bot:  [resposta sobre o pedido — ignorou estoque]",
    "",
    "Você: quero CANCELAR pedido FF-2025-0001",
    "Bot:  Para consultar o STATUS do pedido...",
], [
    "Você: pedido FF-2025-0001",
    "Você: tem algodão em estoque?",
    "Bot:  Sim, temos algodão básico em estoque: 4200m disponíveis. ✓",
    "",
    "Você: quero CANCELAR pedido FF-2025-0001",
    "Bot:  Para cancelar o pedido FF-2025-0001, fale com vendas... ✓",
])

px = p(doc)
add_run(px, "O que mudou: ", bold=True)
px.add_run("o classifier agora roda regras por VERBO primeiro:")

code_block(doc, [
    "# Regras por verbo (alta prioridade)",
    "if re.search(r'\\bcancelar\\b', t):",
    "    return 'cancelar_pedido'",
    "if re.search(r'\\bestoque\\b|tem (algodao|dry.?fit|...)', t):",
    "    return 'disponibilidade_materiais'",
    "",
    "# Só depois checa slot:",
    "if slots_turno.get('numero_pedido'):     # APENAS slots_turno",
    "    return 'status_pedido'",
])

px = p(doc)
add_run(px, "Detalhe importante: ", bold=True)
px.add_run("a regra agora checa ")
add_run(px, "slots_turno", mono=True)
px.add_run(" (o que o usuário disse AGORA), não ")
add_run(px, "slots_efetivos", mono=True)
px.add_run(" (que inclui herança da sessão). Isso evita que o numero_pedido herdado de turnos "
           "anteriores domine a conversa toda. Também criei duas intenções novas no responder: ")
add_run(px, "cancelar_pedido", mono=True)
px.add_run(" e ")
add_run(px, "disponibilidade_materiais", mono=True)
px.add_run(", que consultam os CSVs apropriados.")

# ===== BUG 7 =====
h(doc, "3.6  CRÍTICO 7 — Gradio: cada usuário tem sua própria sessão", level=2, color=ROXO)

p(doc, "Esse bug era assustador: na demo com 2 abas abertas, um usuário enxergava o orçamento do "
       "outro porque a sessão era uma variável GLOBAL no módulo.")

antes_depois(doc, [
    "# front_gradio.py (antes)",
    "sessao = {                       # global ← compartilhada!",
    "    'ativa': False,",
    "    'ultimo_assunto': None,",
    "    'slots_acumulados': {},",
    "}",
    "",
    "def chat(mensagem, historico):",
    "    slots = extrair_slots(mensagem, sessao)   # usa a global",
    "    ...",
], [
    "# front_gradio.py (depois)",
    "sessao_state = gr.State(value=None)   # estado POR sessão",
    "",
    "def processar(mensagem, sessao):",
    "    if sessao is None:",
    "        sessao = criar_sessao()",
    "    ...",
])

px = p(doc)
add_run(px, "Verificação: ", bold=True)
px.add_run("rodei dois usuários simultâneos (A pediu 200 polos com bordado, B pediu 100 moletons). "
           "Depois cada um perguntou 'e o prazo?'. Resultado: cada um recebeu o prazo do SEU "
           "próprio pedido, sem vazamento. ")
add_run(px, "Bug 7 morto.", bold=True, color=DEPOIS)

# ===== BUG 22 =====
h(doc, "3.7  CRÍTICO 22 — 'sem bordado' agora é entendido como negação", level=2, color=ROXO)

antes_depois(doc, [
    "Você: preço de 100 polos SEM bordado",
    "Bot:  Para 100 polos COM BORDADO: R$ 88/peça  ← ignorou o 'sem'",
], [
    "Você: preço de 100 polos SEM bordado",
    "Bot:  Para 100 polos: R$ 71/peça  ← sem bordado, certinho",
])

px = p(doc)
add_run(px, "O que mudou: ", bold=True)
px.add_run("o extractor agora procura primeiro o padrão de negação. Se acha, marca personalizacao='nenhuma' "
           "e nem chega a procurar 'bordado'/'silk'/etc no resto da frase:")

code_block(doc, [
    "# detecta 'sem bordado', 'sem silk', 'nenhuma personalização'",
    "neg_pers = re.search(",
    "    r'(?:sem|nao quero|sem nenhuma?)\\s+'",
    "    r'(bordad[oa]s?|silk\\w*|serigrafia|dtf|estamp\\w*|...)',",
    "    t",
    ")",
    "if neg_pers:",
    "    slots['personalizacao'] = 'nenhuma'",
])

# ===== EXTRAS =====
h(doc, "3.8  Bônus: outros bugs consertados de quebra", level=2, color=AZUL)

p(doc, "Enquanto eu mexia no extractor/classifier, aproveitei pra consertar mais coisas que "
       "estavam na auditoria:")

doc.add_paragraph(
    "MÉDIO 13 — Plurais 'vestidos' e 'calças' agora detectam produto. Adicionei essas chaves "
    "ao dicionário com regex 's?' pra cobrir singular e plural ao mesmo tempo.",
    style="List Bullet")

doc.add_paragraph(
    "MÉDIO 16 — Letra 'M' como tamanho agora exige prefixo 'tamanho M' ou 'grade M'. Antes "
    "qualquer 'm' isolada virava grade=adulto.",
    style="List Bullet")

doc.add_paragraph(
    "MÉDIO 17 — Slot 'urgente' agora aparece na resposta com taxa de 20-40%. Antes era "
    "extraído mas ignorado.",
    style="List Bullet")

doc.add_paragraph(
    "MÉDIO 18 — Produto agora é o MAIS À ESQUERDA na frase, não o último do dicionário. "
    "Resolve 'a regata fica boa com legging?' que dava produto=legging.",
    style="List Bullet")

doc.add_paragraph(
    "MÉDIO 26 — Resposta agora respeita singular/plural ('1 peça' vs 'N peças').",
    style="List Bullet")

doc.add_paragraph(
    "MÉDIO 29 — 'status_pedido' agora consulta lookup_pedidos.csv pra listar as etapas, "
    "em vez de só dizer 'fale com logística'.",
    style="List Bullet")

doc.add_paragraph(
    "ALTO 8 — Conflito quantidade/prazo: agora prazo é extraído primeiro e o trecho é removido "
    "antes de buscar quantidade. 'preciso em 100 dias' não vira mais quantidade=100.",
    style="List Bullet")

doc.add_paragraph(
    "ALTO 10 — 'selecao_opcao' não vira 'ultimo_assunto'. Quando você responde a um menu, "
    "o bot preserva o assunto real anterior pra herança seguinte funcionar.",
    style="List Bullet")

doc.add_paragraph(
    "ALTO 11 — Viabilidade vence qtd_grande_volume quando há prazo+quantidade.",
    style="List Bullet")

doc.add_paragraph(
    "ALTO 30 — Pergunta clarificadora: se você pede preço/prazo sem informar produto, "
    "o bot agora pergunta em vez de mandar pra 'setor de vendas'.",
    style="List Bullet")

doc.add_paragraph(
    "DESIGN 21 — extrair_slots virou função PURA. Não toca em sessão. Quem mescla com contexto "
    "é a função merge_com_contexto, separadamente. Isso resolve uma raiz arquitetural de vários "
    "bugs.",
    style="List Bullet")

# ---------- ARQUITETURA NOVA ----------
doc.add_page_break()
h(doc, "4. Como o pipeline funciona agora", level=1)

p(doc, "Cada turno do usuário passa por 5 etapas claras. Quem chamar precisa fazer nessa ordem:")

code_block(doc, [
    "# 1) extrai do turno atual (puro — não toca em sessão)",
    "em_menu = bool(sessao.get('aguardando_opcao'))",
    "slots_turno = extrair_slots(mensagem, em_menu=em_menu)",
    "",
    "# 2) mescla com contexto, aplicando invalidação",
    "slots_efetivos = merge_com_contexto(slots_turno, sessao)",
    "",
    "# 3) classifica (precisa dos dois: turno + efetivos)",
    "intencao = classificar(mensagem, slots_turno, slots_efetivos,",
    "                       dados['intencoes'], sessao)",
    "",
    "# 4) responde",
    "resposta = responder(intencao, slots_efetivos, dados, sessao, mensagem)",
    "",
    "# 5) persiste mudanças na sessão",
    "atualizar_sessao_pos_turno(sessao, mensagem, slots_efetivos, intencao, resposta)",
])

p(doc, "Essa separação por etapas é o que permite o classifier saber a diferença entre "
       "'o usuário disse AGORA' e 'estamos arrastando do contexto'. Sem isso, regras antigas "
       "viravam booby-trap (ex: numero_pedido herdado fazia todo turno virar status_pedido).")

# ---------- RESULTADOS ----------
h(doc, "5. Verificação: bateria dos 15 cenários", level=1)
p(doc, "Todos os 15 cenários da auditoria passam agora. Os mais importantes:")

tab = doc.add_table(rows=16, cols=3)
tab.style = "Light Grid Accent 1"
hdr = ["#", "Cenário", "Status"]
for j, lbl in enumerate(hdr):
    tab.rows[0].cells[j].text = ""
    add_run(tab.rows[0].cells[j].paragraphs[0], lbl, bold=True)

cenarios = [
    ("1", "Garçom (preço → e com bordado? → e o prazo?)", "OK"),
    ("2", "Mudar produto não gruda bordado", "OK"),
    ("3", "Menu não corrompe quantidade", "OK"),
    ("4", "Despedida no meio não reseta", "OK"),
    ("5", "Número de pedido não vira quantidade", "OK"),
    ("6", "Cancelar funciona", "OK"),
    ("7", "Estoque depois de pedido", "OK"),
    ("8", "Follow-up sem número", "OK"),
    ("9", "Viabilidade vence qtd_grande_volume", "OK"),
    ("10", "Plural 'vestidos'", "OK"),
    ("11", "Plural 'calças jeans'", "OK"),
    ("12", "Conflito quantidade/prazo → pergunta clarificadora", "OK"),
    ("13", "Ofensa + retomada", "OK"),
    ("14", "Casual no meio mantém contexto", "OK"),
    ("15", "Urgência com aviso de taxa", "OK"),
]
for i, (n, desc, st) in enumerate(cenarios, 1):
    tab.rows[i].cells[0].text = n
    tab.rows[i].cells[1].text = desc
    tab.rows[i].cells[2].text = st

p(doc)
nota = p(doc)
add_run(nota, "Bug 7 (Gradio) ", bold=True)
nota.add_run("verificado separadamente com duas sessões simultâneas — cada usuário manteve sua "
             "própria conversa, sem vazamento.")

# ---------- DEMO ----------
h(doc, "6. Como rodar pra demonstrar", level=1)

px = p(doc)
add_run(px, "Terminal: ", bold=True)
add_run(px, "python main.py", mono=True)
px.add_run(" (no terminal você pode digitar ")
add_run(px, "/contexto", mono=True)
px.add_run(" pra ver o estado atual da memória — ótimo pra mostrar pra professora)")

px = p(doc)
add_run(px, "Interface Gradio: ", bold=True)
add_run(px, "python front_gradio.py", mono=True)
px.add_run(" (abre o navegador. Tem um painel lateral mostrando a 'memória da conversa' "
           "atualizando em tempo real)")

px = p(doc)
add_run(px, "API REST: ", bold=True)
add_run(px, "uvicorn app:app --reload", mono=True)
px.add_run(" (sobe em http://localhost:8000, docs em /docs)")

# ---------- ARQUIVOS ----------
h(doc, "7. Quais arquivos mudaram", level=1)

tab = doc.add_table(rows=8, cols=2)
tab.style = "Light Grid Accent 1"
hdr = ["Arquivo", "O que mudou"]
for j, lbl in enumerate(hdr):
    tab.rows[0].cells[j].text = ""
    add_run(tab.rows[0].cells[j].paragraphs[0], lbl, bold=True)

arquivos = [
    ("bot/contexto.py",
     "REESCRITO. Novo design: foco_atual + historico_turnos. "
     "Funções merge_com_contexto, atualizar_sessao_pos_turno. "
     "is_despedida restritivo, is_casual expandido."),
    ("bot/extractor.py",
     "REESCRITO. Função PURA. Extrai numero_pedido e prazo_desejado primeiro. "
     "Detecção de negação. Plurais. Produto mais à esquerda. Param 'em_menu'."),
    ("bot/classifier.py",
     "REESCRITO. Recebe slots_turno + slots_efetivos. Regras por verbo primeiro. "
     "Nova intenção 'prazo_sem_contexto'. Herança melhor pra follow-ups."),
    ("bot/responder.py",
     "REESCRITO. Singular/plural. Novos handlers cancelar_pedido, disponibilidade_materiais, "
     "prazo_sem_contexto. Aviso de urgência. status_pedido usa lookup_pedidos."),
    ("main.py",
     "REESCRITO. Pipeline novo de 5 etapas. Comando /contexto pra debug. UTF-8 no Windows."),
    ("app.py",
     "REESCRITO. Mesmo pipeline. Novo endpoint GET /sessao/{id} pra debug/demo."),
    ("front_gradio.py",
     "REESCRITO. gr.State() por usuário (CRÍTICO 7). Painel lateral de memória. "
     "Botão 'limpar conversa'. Tema atualizado. API Gradio 6 corrigida."),
]
for i, (k, v) in enumerate(arquivos, 1):
    tab.rows[i].cells[0].text = k
    tab.rows[i].cells[1].text = v

# ---------- PRÓXIMOS PASSOS ----------
h(doc, "8. Próximos passos sugeridos", level=1)
doc.add_paragraph(
    "Testar a interface real abrindo o Gradio em duas abas pra confirmar isolamento.",
    style="List Number")
doc.add_paragraph(
    "Apresentar pra professora mostrando o painel de 'memória da conversa' atualizando — "
    "é a prova visual do 'Poder do Contexto' do PDF.",
    style="List Number")
doc.add_paragraph(
    "Se sobrar tempo: Bloco B (follow-up sem número, refator de extractor.py em testes "
    "automatizados, perguntas clarificadoras pra 'qual produto?').",
    style="List Number")
doc.add_paragraph(
    "Atualizar o lookup_estoque_materiais.csv: as datas de reposição estão em 2025 "
    "(passadas).",
    style="List Number")
doc.add_paragraph(
    "Commitar tudo. Mensagem de commit sugerida: 'fix(bot): bloco A da auditoria — "
    "contexto, negação, menus, cancelar, gradio gr.State'.",
    style="List Number")

p(doc)
fim = p(doc)
fim.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(fim, "— fim do documento —", italic=True, color=CINZA)

doc.save(OUT)
print(f"OK: {OUT}")
print(f"Tamanho: {OUT.stat().st_size} bytes")
