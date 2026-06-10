"""
Gera estudo_python.docx — Material didático completo pra Bia estudar pra prova.

Cobre:
1. Fundamentos de Python (do zero)
2. Conceitos do projeto (regex, dataframes, funções puras)
3. Tabela de tradução Python ⇄ Portugol
4. Código do bot linha por linha
5. Questões prováveis de prova com gabarito

Execute: python docs/gerar_doc_estudo.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

OUT = Path(__file__).parent / "estudo_python.docx"

# Cores
ROXO   = RGBColor(0x76, 0x4A, 0xB0)
AZUL   = RGBColor(0x21, 0x65, 0xA3)
VERDE  = RGBColor(0x16, 0xA3, 0x4A)
LARANJA = RGBColor(0xE0, 0x7B, 0x00)
CINZA  = RGBColor(0x55, 0x55, 0x55)
PRETO  = RGBColor(0x00, 0x00, 0x00)
VERM   = RGBColor(0xC0, 0x39, 0x2B)


def add_run(p, text, bold=False, italic=False, color=None, mono=False, size=None, underline=False):
    r = p.add_run(text)
    if bold: r.bold = True
    if italic: r.italic = True
    if underline: r.underline = True
    if color: r.font.color.rgb = color
    r.font.name = "Consolas" if mono else "Calibri"
    if mono and size is None: size = 9.5
    if size: r.font.size = Pt(size)
    return r


def h(doc, text, level=1, color=None):
    p = doc.add_heading(level=level)
    r = p.add_run(text)
    if color: r.font.color.rgb = color
    return p


def p(doc, text=""):
    return doc.add_paragraph(text)


def code(doc, lines, cor="F4F4F4"):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(0.5)
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    add_run(para, lines if isinstance(lines, str) else "\n".join(lines),
            mono=True, color=CINZA, size=9.5)
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), cor)
    pPr.append(shd)


def py(doc, lines):
    """Bloco de código Python (fundo cinza claro)."""
    code(doc, lines, "F4F4F4")


def portugol(doc, lines):
    """Bloco de portugol (fundo amarelo claro)."""
    code(doc, lines, "FFF8DC")


def nota(doc, texto):
    """Nota destacada (fundo azul claro)."""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(0.5)
    add_run(para, "💡 ", bold=True)
    add_run(para, texto, italic=True, color=CINZA)
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'EAF4FB')
    pPr.append(shd)


def atencao(doc, texto):
    """Aviso (fundo amarelo)."""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(0.5)
    add_run(para, "⚠️ ", bold=True)
    add_run(para, texto, italic=True, color=CINZA)
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'FFF4D6')
    pPr.append(shd)


# ============================================================
doc = Document()
for s in doc.sections:
    s.left_margin = Cm(2.0); s.right_margin = Cm(2.0)
    s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)

# ============== CAPA ==============
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(t, "\n\n\n", size=14)
add_run(t, "Estudando Python\n", bold=True, size=32, color=ROXO)
add_run(t, "com o Fashion Flow Bot\n\n", bold=True, size=20, color=AZUL)

s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(s, "Material de estudo pra prova\n", italic=True, size=14, color=CINZA)
add_run(s, "Fundamentos + código real + portugol + simulado\n", italic=True, color=CINZA)
add_run(s, "\nBia Costa  •  Engenharia 1º sem  •  2026-06-10\n", size=10, color=CINZA)

p(doc, "")
caixa = p(doc)
caixa.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(caixa, "📚 Use esse doc como apostila — leia uma seção por vez, faça os exercícios, "
              "e tente reescrever os trechos de código de cabeça pra fixar.",
        italic=True, color=CINZA, size=11)

doc.add_page_break()

# ============== SUMÁRIO ==============
h(doc, "Como ler esse documento", level=1)
p(doc, "O doc está dividido em 5 partes — leia na ordem na primeira passada, depois consulta direto na parte que precisar.")

doc.add_paragraph("Parte I — Fundamentos do Python (do zero)", style="List Number")
doc.add_paragraph("Parte II — Conceitos do projeto que aparecem no código", style="List Number")
doc.add_paragraph("Parte III — Tradução Python ⇄ Portugol (tabela de equivalência)", style="List Number")
doc.add_paragraph("Parte IV — O código do bot linha por linha", style="List Number")
doc.add_paragraph("Parte V — Simulado: questões prováveis de prova com gabarito", style="List Number")

doc.add_page_break()


# ##################################################################
# PARTE I — FUNDAMENTOS DO PYTHON
# ##################################################################
h(doc, "Parte I — Fundamentos do Python", level=1, color=ROXO)
p(doc, "Nessa parte cobrimos tudo o que você precisa saber pra entender o resto do doc — e pra resolver a prova.")

# ── 1. O que é um programa Python ──
h(doc, "1.1 O que é um programa Python?", level=2)
p(doc, "Um programa Python é um arquivo de texto com extensão .py. Quando você 'roda' esse arquivo, "
       "o computador lê linha por linha, de cima pra baixo, e vai executando o que está escrito. "
       "Diferente de Java ou C, NÃO existe compilação manual — é só rodar com 'python arquivo.py'.")

py(doc, [
    "# arquivo: ola.py",
    'print("Olá, mundo!")',
    'print("Eu sei Python.")',
])

p(doc, "Quando você roda 'python ola.py', aparece no terminal:")
code(doc, ["Olá, mundo!", "Eu sei Python."], "FFFFFF")

nota(doc, "O sinal # serve pra comentário — Python ignora tudo depois dele na linha. "
          "Use pra anotar o que o código faz.")

# ── 2. Variáveis ──
h(doc, "1.2 Variáveis e atribuição", level=2)
p(doc, "Variável é uma 'caixa' onde a gente guarda um valor pra usar depois. "
       "Em Python, criar variável é só escrever nome + igual + valor:")
py(doc, [
    "nome = 'Bia'",
    "idade = 19",
    "altura = 1.65",
    "estudante = True",
])

p(doc, "Em portugol seria:")
portugol(doc, [
    "var",
    "   nome: caractere",
    "   idade: inteiro",
    "   altura: real",
    "   estudante: logico",
    "inicio",
    '   nome <- "Bia"',
    "   idade <- 19",
    "   altura <- 1.65",
    "   estudante <- verdadeiro",
])

atencao(doc, "Diferença grande: em Python NÃO declaramos o tipo da variável. O Python descobre sozinho. "
              "Em portugol/VisualG você precisa declarar 'idade: inteiro' antes de usar.")

# ── 3. Tipos básicos ──
h(doc, "1.3 Os tipos básicos", level=2)

tab = doc.add_table(rows=6, cols=4)
tab.style = "Light Grid Accent 1"
hdr = ["Tipo Python", "Português", "Exemplo Python", "Portugol"]
for j, lbl in enumerate(hdr):
    tab.rows[0].cells[j].text = ""
    add_run(tab.rows[0].cells[j].paragraphs[0], lbl, bold=True)

dados = [
    ("int",   "Número inteiro",       "x = 42",                "inteiro"),
    ("float", "Número decimal",       "preco = 19.90",         "real"),
    ("str",   "Texto (string)",       'nome = "Maria"',        "caractere"),
    ("bool",  "Verdadeiro ou falso",  "ativo = True",          "logico"),
    ("None",  "'Nada'/vazio",         "valor = None",          "(não existe equivalente direto)"),
]
for i, row in enumerate(dados, 1):
    for j, val in enumerate(row):
        tab.rows[i].cells[j].text = val

p(doc)
p(doc, "Você pode descobrir o tipo de uma variável com type():")
py(doc, [
    "x = 10",
    "print(type(x))   # <class 'int'>",
    "",
    "y = 3.14",
    "print(type(y))   # <class 'float'>",
])

# ── 4. Strings ──
h(doc, "1.4 Strings (textos)", level=2)
p(doc, "String é qualquer texto entre aspas. Pode ser aspas simples ou duplas — tanto faz, desde que feche igual abriu.")
py(doc, [
    "a = 'oi'",
    'b = "tudo bem?"',
    "c = 'pode ter \"aspas duplas dentro\" usando simples por fora'",
])

p(doc)
add_run(p(doc), "F-string — a forma mais usada de formatar texto:", bold=True)
py(doc, [
    "nome = 'Bia'",
    "idade = 19",
    "print(f'Oi {nome}, você tem {idade} anos.')",
    "# saída: Oi Bia, você tem 19 anos.",
])

p(doc, "Dentro do f'...' tudo entre { } é AVALIADO como código Python. Pode fazer conta:")
py(doc, [
    "preco_unit = 67.41",
    "qtd = 200",
    "print(f'Total: R$ {preco_unit * qtd:.2f}')",
    "# saída: Total: R$ 13482.00",
])

nota(doc, "O :.2f depois da variável significa 'formate com 2 casas decimais'. "
          "Sem isso, 13482.0 apareceria 13482.0 mesmo.")

p(doc)
add_run(p(doc), "Métodos úteis de string:", bold=True)
py(doc, [
    "msg = '  Olá, Mundo!  '",
    "print(msg.lower())       # '  olá, mundo!  '",
    "print(msg.upper())       # '  OLÁ, MUNDO!  '",
    "print(msg.strip())       # 'Olá, Mundo!'  (tira espaços nas pontas)",
    "print(msg.replace('Olá', 'Oi'))  # '  Oi, Mundo!  '",
    "print('mundo' in msg.lower())    # True",
    "print(len(msg))                  # 17 (quantidade de caracteres)",
])

# ── 5. Operadores ──
h(doc, "1.5 Operadores matemáticos e de comparação", level=2)
tab = doc.add_table(rows=11, cols=3)
tab.style = "Light Grid Accent 1"
for j, lbl in enumerate(["Python", "Significado", "Exemplo"]):
    tab.rows[0].cells[j].text = ""
    add_run(tab.rows[0].cells[j].paragraphs[0], lbl, bold=True)

ops = [
    ("+",  "soma",                 "3 + 2 → 5"),
    ("-",  "subtração",            "5 - 1 → 4"),
    ("*",  "multiplicação",        "4 * 3 → 12"),
    ("/",  "divisão (sempre float)","10 / 4 → 2.5"),
    ("//", "divisão inteira",       "10 // 4 → 2"),
    ("%",  "resto da divisão",      "10 % 3 → 1"),
    ("**", "potência",              "2 ** 8 → 256"),
    ("==", "igual a (compara)",     "x == 5"),
    ("!=", "diferente de",          "x != 5"),
    ("<", ">", "<=", ">="),
]
# refazendo último: vou colocar comparações em outra linha
ops = [
    ("+",  "soma",                 "3 + 2 → 5"),
    ("-",  "subtração",            "5 - 1 → 4"),
    ("*",  "multiplicação",        "4 * 3 → 12"),
    ("/",  "divisão (sempre float)","10 / 4 → 2.5"),
    ("//", "divisão inteira",       "10 // 4 → 2"),
    ("%",  "resto da divisão",      "10 % 3 → 1"),
    ("**", "potência",              "2 ** 8 → 256"),
    ("==", "igual a (compara)",     "x == 5  → True/False"),
    ("!=", "diferente de",          "x != 5"),
    ("<, >, <=, >=", "comparações", "x >= 18"),
]
for i, row in enumerate(ops, 1):
    for j, val in enumerate(row):
        tab.rows[i].cells[j].text = val

p(doc)
atencao(doc, "Não confunda = com ==. '=' atribui (x = 5 → guarda 5 em x). "
              "'==' compara (x == 5 → True se x vale 5).")

p(doc)
add_run(p(doc), "Operadores lógicos:", bold=True)
py(doc, [
    "idade = 20",
    "tem_doc = True",
    "",
    "# E lógico",
    "if idade >= 18 and tem_doc:",
    "    print('pode entrar')",
    "",
    "# OU lógico",
    "if idade < 5 or idade > 60:",
    "    print('preço especial')",
    "",
    "# NÃO lógico",
    "if not tem_doc:",
    "    print('falta documento')",
])

# ── 6. Listas ──
h(doc, "1.6 Listas (sequências)", level=2)
p(doc, "Lista é uma sequência ordenada de valores. Em Python usa colchetes [ ].")
py(doc, [
    "frutas = ['maçã', 'banana', 'uva']",
    "numeros = [10, 20, 30, 40]",
    "misturada = [1, 'oi', True, 3.14]   # pode misturar tipos",
    "vazia = []",
])

p(doc, "Acessar elementos (começa em 0!):")
py(doc, [
    "frutas = ['maçã', 'banana', 'uva']",
    "print(frutas[0])    # 'maçã'  ← primeiro",
    "print(frutas[1])    # 'banana'",
    "print(frutas[-1])   # 'uva'   ← último (índice negativo conta do fim)",
    "print(len(frutas))  # 3 (tamanho)",
])

p(doc, "Adicionar e remover:")
py(doc, [
    "frutas.append('laranja')   # adiciona no fim",
    "frutas.remove('banana')    # remove a primeira ocorrência",
    "frutas.pop()               # remove e devolve o último",
])

# ── 7. Dicionários ──
h(doc, "1.7 Dicionários (chave → valor)", level=2)
p(doc, "Dicionário é uma 'agenda': pra cada CHAVE você guarda um VALOR. Em Python usa chaves { }.")
py(doc, [
    "aluno = {",
    "    'nome': 'Bia',",
    "    'idade': 19,",
    "    'curso': 'Engenharia',",
    "}",
])

p(doc, "Acessar valores:")
py(doc, [
    "print(aluno['nome'])           # 'Bia'",
    "print(aluno.get('idade'))      # 19  (forma 'segura')",
    "print(aluno.get('cep'))        # None  (não dá erro se não existir)",
    "print(aluno.get('cep', '?'))   # '?'   (com default)",
])

p(doc, "Adicionar/mudar valores:")
py(doc, [
    "aluno['email'] = 'bia@exemplo.com'   # adiciona",
    "aluno['idade'] = 20                   # muda o existente",
    "del aluno['curso']                    # remove",
])

p(doc, "Percorrer um dicionário:")
py(doc, [
    "for chave, valor in aluno.items():",
    "    print(f'{chave}: {valor}')",
])

nota(doc, "Dicionário é o tipo MAIS USADO no nosso bot. A sessão é um dicionário, "
          "os slots são um dicionário, os CSVs viram lista de dicionários. Domine isso e domina o código.")

# ── 8. if/elif/else ──
h(doc, "1.8 Decisões: if / elif / else", level=2)
py(doc, [
    "idade = 19",
    "",
    "if idade < 18:",
    "    print('menor de idade')",
    "elif idade < 65:",
    "    print('adulto')",
    "else:",
    "    print('idoso')",
])

p(doc, "Em portugol:")
portugol(doc, [
    "se idade < 18 entao",
    '   escreva("menor de idade")',
    "senao se idade < 65 entao",
    '   escreva("adulto")',
    "senao",
    '   escreva("idoso")',
    "fimse",
])

atencao(doc, "Em Python, a INDENTAÇÃO (espaços antes da linha) define o que pertence ao if. "
              "NÃO existem 'fimse' ou chaves. Use sempre 4 espaços por nível.")

py(doc, [
    "# CERTO",
    "if idade >= 18:",
    "    print('adulto')",
    "    print('pode votar')   # ainda dentro do if",
    "print('fim')               # FORA do if",
    "",
    "# ERRADO — IndentationError",
    "if idade >= 18:",
    "print('adulto')   # sem indentação, Python não entende",
])

# ── 9. for ──
h(doc, "1.9 Loops: for", level=2)
p(doc, "Use for quando você sabe quantas vezes vai repetir, ou quando quer percorrer uma sequência.")
py(doc, [
    "# percorrer lista",
    "for fruta in ['maçã', 'banana', 'uva']:",
    "    print(fruta)",
    "",
    "# percorrer números",
    "for i in range(5):           # 0, 1, 2, 3, 4",
    "    print(i)",
    "",
    "for i in range(1, 6):        # 1, 2, 3, 4, 5",
    "    print(i)",
    "",
    "for i in range(0, 10, 2):    # 0, 2, 4, 6, 8 (pula de 2 em 2)",
    "    print(i)",
])

p(doc, "Portugol equivalente:")
portugol(doc, [
    "para i de 1 ate 5 faca",
    "   escreva(i)",
    "fimpara",
])

# ── 10. while ──
h(doc, "1.10 Loops: while", level=2)
p(doc, "while repete ENQUANTO uma condição for verdadeira. Usado quando não sabemos quantas vezes.")
py(doc, [
    "contador = 0",
    "while contador < 5:",
    "    print(contador)",
    "    contador = contador + 1   # ou: contador += 1",
])

p(doc, "Loop infinito (cuidado!):")
py(doc, [
    "while True:",
    "    msg = input('Digite: ')",
    "    if msg == 'sair':",
    "        break    # quebra o loop",
    "    print(f'Você digitou: {msg}')",
])

p(doc, "Portugol:")
portugol(doc, [
    "contador <- 0",
    "enquanto contador < 5 faca",
    "   escreva(contador)",
    "   contador <- contador + 1",
    "fimenquanto",
])

# ── 11. Funções ──
h(doc, "1.11 Funções (def / return)", level=2)
p(doc, "Função é um 'pedaço de código' que tem nome e pode ser reusado. Define com 'def', executa chamando o nome.")
py(doc, [
    "def saudar(nome):",
    "    return f'Olá, {nome}!'",
    "",
    "mensagem = saudar('Bia')",
    "print(mensagem)   # Olá, Bia!",
])

p(doc, "Funções podem ter múltiplos parâmetros e valores padrão:")
py(doc, [
    "def calcular_total(qtd, preco_unit, desconto=0):",
    "    subtotal = qtd * preco_unit",
    "    return subtotal * (1 - desconto / 100)",
    "",
    "print(calcular_total(10, 50))            # 500 (desconto 0 por padrão)",
    "print(calcular_total(10, 50, 20))        # 400 (com 20% de desconto)",
    "print(calcular_total(qtd=10, preco_unit=50, desconto=20))  # mesma coisa",
])

p(doc, "Funções podem retornar vários valores (vira tupla):")
py(doc, [
    "def dividir(a, b):",
    "    quociente = a // b",
    "    resto = a % b",
    "    return quociente, resto",
    "",
    "q, r = dividir(10, 3)",
    "print(q, r)   # 3 1",
])

p(doc, "Portugol equivalente:")
portugol(doc, [
    "funcao saudar(nome: caractere): caractere",
    "inicio",
    '   retorne "Olá, " + nome + "!"',
    "fimfuncao",
])

nota(doc, "Funções 'puras' são as que SÓ usam seus parâmetros e devolvem um valor — não tocam em variáveis "
          "de fora. Nosso extrair_slots virou pura na correção do projeto. Funções puras são mais fáceis de testar.")

# ── 12. try / except ──
h(doc, "1.12 Tratamento de erros: try / except", level=2)
p(doc, "Quando algo pode dar erro (entrada inválida do usuário, arquivo não encontrado, etc), "
       "use try/except pra não derrubar o programa.")
py(doc, [
    "try:",
    "    numero = int(input('Digite um número: '))",
    "    print(f'Você digitou {numero}')",
    "except ValueError:",
    "    print('Isso não é um número válido!')",
])

p(doc, "Vários tipos de erro:")
py(doc, [
    "try:",
    "    abrir_arquivo()",
    "except FileNotFoundError:",
    "    print('Arquivo não encontrado')",
    "except PermissionError:",
    "    print('Sem permissão')",
    "except Exception as e:    # qualquer outro erro",
    "    print(f'Deu ruim: {e}')",
])

# ── 13. Imports ──
h(doc, "1.13 Importando bibliotecas", level=2)
p(doc, "Pra usar código de outros arquivos ou bibliotecas externas, usa 'import'.")
py(doc, [
    "import math                     # pega tudo de math",
    "print(math.pi)                  # 3.141592...",
    "print(math.sqrt(16))            # 4.0",
    "",
    "from math import pi, sqrt       # pega só o que precisa",
    "print(pi)                       # sem precisar do 'math.'",
    "print(sqrt(16))",
    "",
    "import pandas as pd             # com apelido",
    "df = pd.read_csv('dados.csv')",
])

doc.add_page_break()


# ##################################################################
# PARTE II — CONCEITOS DO PROJETO
# ##################################################################
h(doc, "Parte II — Conceitos do projeto", level=1, color=ROXO)
p(doc, "Coisas específicas que aparecem no código do bot e merecem uma explicação separada.")

# ── 2.1 Regex ──
h(doc, "2.1 Regex (expressões regulares)", level=2)
p(doc, "Regex é um 'padrão' pra buscar texto. Tipo achar todos os números numa frase, "
       "ou descobrir se a mensagem tem 'bordado'. No nosso bot, regex é usado em CADA arquivo do bot/.")

py(doc, [
    "import re",
    "",
    "mensagem = 'quero 100 polos com bordado'",
    "",
    "# Procurar um padrão",
    "match = re.search(r'(\\d+)', mensagem)",
    "if match:",
    "    print(match.group())   # '100'",
    "    print(match.group(1))  # '100' (primeiro grupo)",
])

p(doc, "Símbolos importantes:")
tab = doc.add_table(rows=10, cols=3)
tab.style = "Light Grid Accent 1"
for j, lbl in enumerate(["Símbolo", "Significado", "Exemplo"]):
    tab.rows[0].cells[j].text = ""
    add_run(tab.rows[0].cells[j].paragraphs[0], lbl, bold=True)

simbolos = [
    (r"\d", "um dígito (0-9)", r"\d → casa '5'"),
    (r"\d+", "um ou mais dígitos", r"\d+ → casa '100'"),
    (r"\w", "letra, número ou _", r"\w+ → casa 'abc'"),
    (r"\s", "espaço, tab, quebra de linha", r"\s+ → casa '   '"),
    (r"\b", "borda de palavra", r"\bola\b → 'ola' isolada"),
    (".", "qualquer caractere", "a.b → 'aXb', 'a@b'"),
    ("?", "0 ou 1 vez", "ab?c → 'ac' ou 'abc'"),
    ("|", "ou", "(sim|nao) → 'sim' OU 'nao'"),
    ("[...]", "qualquer um dos caracteres", "[aeiou] → vogal"),
]
for i, row in enumerate(simbolos, 1):
    for j, val in enumerate(row):
        tab.rows[i].cells[j].text = val

p(doc)
add_run(p(doc), "Exemplo do nosso bot:", bold=True)
py(doc, [
    r"# extrai número de pedido tipo FF-2024-0123",
    r"match = re.search(r'FF-\d{4}-\d{4}', mensagem)",
    r"# \d{4} = exatamente 4 dígitos",
])

# ── 2.2 List comprehension ──
h(doc, "2.2 List comprehension (criar lista em uma linha)", level=2)
p(doc, "Forma compacta de criar uma lista a partir de outra.")
py(doc, [
    "# Forma longa",
    "dobrados = []",
    "for x in [1, 2, 3, 4]:",
    "    dobrados.append(x * 2)",
    "print(dobrados)   # [2, 4, 6, 8]",
    "",
    "# Forma com list comprehension",
    "dobrados = [x * 2 for x in [1, 2, 3, 4]]",
    "print(dobrados)   # [2, 4, 6, 8]",
])

p(doc, "Com filtro:")
py(doc, [
    "# só os pares dobrados",
    "pares_dobrados = [x * 2 for x in range(10) if x % 2 == 0]",
    "print(pares_dobrados)   # [0, 4, 8, 12, 16]",
])

# ── 2.3 Dict comprehension ──
h(doc, "2.3 Dict comprehension", level=2)
py(doc, [
    "# transformar a lista de produtos em dicionário",
    "produtos = ['camiseta', 'polo', 'moletom']",
    "precos = {p: 30 for p in produtos}",
    "print(precos)   # {'camiseta': 30, 'polo': 30, 'moletom': 30}",
])

# ── 2.4 Função pura vs side effect ──
h(doc, "2.4 Função pura vs efeito colateral", level=2)
p(doc, "Função PURA: só usa os parâmetros e devolve um valor — não mexe em nada fora.")
py(doc, [
    "# pura: depende só de a e b",
    "def somar(a, b):",
    "    return a + b",
])

p(doc, "Função com EFEITO COLATERAL: mexe em algo de fora.")
py(doc, [
    "contador = 0",
    "",
    "def incrementar():",
    "    global contador     # mexe na variável de fora",
    "    contador += 1",
])

nota(doc, "No nosso bot, o BUG 21 era exatamente esse — a função extrair_slots mexia na sessão "
          "que recebia. Reescrevemos pra ser pura: ela só recebe a mensagem e devolve os slots. "
          "Quem mescla com a sessão é outra função.")

# ── 2.5 Pandas (DataFrame) ──
h(doc, "2.5 Pandas DataFrame — tabelas em Python", level=2)
p(doc, "Pandas é a biblioteca pra trabalhar com tabelas (CSV/Excel) em Python. "
       "Cada tabela vira um DataFrame.")
py(doc, [
    "import pandas as pd",
    "",
    "df = pd.read_csv('lookup_preco.csv')",
    "",
    "# ver as primeiras 5 linhas",
    "print(df.head())",
    "",
    "# filtrar: só linhas onde produto = 'polo'",
    "polos = df[df['produto'] == 'polo']",
    "",
    "# múltiplas condições",
    "filtro = df[",
    "    (df['produto'] == 'polo') &",
    "    (df['qtd_min'] <= 200) &",
    "    (df['qtd_max'] >= 200)",
    "]",
    "",
    "# pegar a primeira linha do filtro",
    "linha = filtro.iloc[0]",
    "print(linha['preco_unitario_estimado'])",
])

# ── 2.6 Decoradores ──
h(doc, "2.6 Decoradores (@) — em FastAPI e Flask", level=2)
p(doc, "Um '@' antes de uma função é um decorador. É uma forma de adicionar comportamento à função "
       "sem mexer nela. No FastAPI, decoradores transformam funções em endpoints HTTP.")
py(doc, [
    "from fastapi import FastAPI",
    "app = FastAPI()",
    "",
    "@app.get('/')                # responde quando alguém pede GET /",
    "def home():",
    "    return {'msg': 'Olá!'}",
    "",
    "@app.post('/chat')           # responde POST /chat",
    "def chat(req):",
    "    # ... processa a mensagem ...",
    "    return {'resposta': '...'}",
])

doc.add_page_break()


# ##################################################################
# PARTE III — TABELA PYTHON ⇄ PORTUGOL
# ##################################################################
h(doc, "Parte III — Tradução Python ⇄ Portugol", level=1, color=ROXO)
p(doc, "Pra prova: o professor pode pedir o algoritmo em Python OU em portugol. Essa tabela traduz "
       "os padrões mais comuns. Decora os fundamentais.")

# atribuição
h(doc, "3.1 Atribuição", level=2)
py(doc, ["nome = 'Bia'", "idade = 19"])
portugol(doc, ['nome <- "Bia"', "idade <- 19"])

# entrada e saída
h(doc, "3.2 Entrada e saída", level=2)
py(doc, [
    "nome = input('Digite seu nome: ')",
    "print('Oi', nome)",
])
portugol(doc, [
    'escreva("Digite seu nome: ")',
    "leia(nome)",
    'escreva("Oi ", nome)',
])

# if
h(doc, "3.3 Decisão", level=2)
py(doc, [
    "if nota >= 7:",
    "    print('aprovado')",
    "elif nota >= 5:",
    "    print('recuperação')",
    "else:",
    "    print('reprovado')",
])
portugol(doc, [
    "se nota >= 7 entao",
    '   escreva("aprovado")',
    "senao se nota >= 5 entao",
    '   escreva("recuperação")',
    "senao",
    '   escreva("reprovado")',
    "fimse",
])

# for
h(doc, "3.4 Repetição contada (for)", level=2)
py(doc, [
    "for i in range(1, 11):",
    "    print(i)",
])
portugol(doc, [
    "para i de 1 ate 10 faca",
    "   escreva(i)",
    "fimpara",
])

# while
h(doc, "3.5 Repetição condicional (while)", level=2)
py(doc, [
    "x = 0",
    "while x < 5:",
    "    print(x)",
    "    x += 1",
])
portugol(doc, [
    "x <- 0",
    "enquanto x < 5 faca",
    "   escreva(x)",
    "   x <- x + 1",
    "fimenquanto",
])

# função
h(doc, "3.6 Função", level=2)
py(doc, [
    "def media(a, b):",
    "    return (a + b) / 2",
    "",
    "resultado = media(7, 9)",
])
portugol(doc, [
    "funcao media(a: real, b: real): real",
    "inicio",
    "   retorne (a + b) / 2",
    "fimfuncao",
])

# vetor / lista
h(doc, "3.7 Vetor (lista)", level=2)
py(doc, [
    "notas = [7.5, 8.0, 6.0]",
    "print(notas[0])     # 7.5",
    "notas.append(9.0)",
    "",
    "for n in notas:",
    "    print(n)",
])
portugol(doc, [
    "var",
    "   notas: vetor[1..4] de real",
    "inicio",
    "   notas[1] <- 7.5",
    "   notas[2] <- 8.0",
    "   notas[3] <- 6.0",
    "   notas[4] <- 9.0",
    "",
    "   para i de 1 ate 4 faca",
    "      escreva(notas[i])",
    "   fimpara",
])

atencao(doc, "Atenção: em Python listas começam em 0; em portugol normalmente começam em 1. "
              "Cuidado nessa hora!")

doc.add_page_break()


# ##################################################################
# PARTE IV — CÓDIGO DO BOT LINHA POR LINHA
# ##################################################################
h(doc, "Parte IV — O código do bot, linha por linha", level=1, color=ROXO)
p(doc, "Aqui vamos passar pelo código do projeto explicando linha por linha. "
       "Vou priorizar os arquivos mais didáticos — main.py, contexto.py e partes do extractor.py.")

# ── loader.py ──
h(doc, "4.1 bot/loader.py — carregar os CSVs", level=2)
p(doc, "Esse arquivo é curtinho, ideal pra começar.")

py(doc, ["import pandas as pd"])
p(doc, "📖 Importa a biblioteca pandas e dá um apelido 'pd' (convenção universal).")

py(doc, ["from pathlib import Path"])
p(doc, "📖 Importa só a classe Path da biblioteca pathlib. Path representa um caminho de arquivo "
       "de forma portável (funciona no Windows, Linux, Mac).")

py(doc, ["DATA_DIR = Path(__file__).parent.parent / 'data'"])
p(doc, "📖 Cria uma constante (nome em maiúsculas por convenção) com o caminho da pasta data/. "
       "__file__ é o caminho do arquivo atual. .parent pega a pasta dele. .parent.parent sobe duas pastas. "
       "Depois junta com '/ data' (a barra ali é operador especial do Path, não divisão).")

py(doc, ["def carregar_dados():"])
p(doc, "📖 Define uma função chamada carregar_dados que não recebe parâmetros.")

py(doc, [
    "    tabelas = {",
    "        'intencoes':                    'intencoes.csv',",
    "        'slots':                        'slots.csv',",
    "        'prazo':                        'lookup_prazo.csv',",
    "        # ... (mais entradas)",
    "    }",
])
p(doc, "📖 Cria um dicionário onde a chave é o NOME que vamos usar e o valor é o NOME DO ARQUIVO. "
       "Isso permite renomear arquivos sem mexer no resto do código.")

py(doc, [
    "    dados = {}",
    "    for nome, arquivo in tabelas.items():",
    "        caminho = DATA_DIR / arquivo",
    "        dados[nome] = pd.read_csv(caminho)",
])
p(doc, "📖 Cria um dict vazio. Pra cada par (nome, arquivo) do dicionário, monta o caminho completo "
       "(DATA_DIR + nome do arquivo) e lê o CSV com pandas. Guarda no dict 'dados' com o nome amigável.")

py(doc, ["    return dados"])
p(doc, "📖 Devolve o dicionário pronto. Quem chamar carregar_dados() vai receber tudo de uma vez.")

# ── main.py ──
h(doc, "4.2 main.py — o ponto de entrada (loop de chat)", level=2)
p(doc, "Esse é o arquivo que você roda com 'python main.py'. É o mais didático.")

py(doc, [
    "import sys",
    "from bot.loader import carregar_dados",
    "from bot.extractor import extrair_slots",
    "from bot.classifier import classificar",
    "from bot.responder import responder",
    "from bot.contexto import (",
    "    criar_sessao, resetar_sessao,",
    "    is_despedida, is_casual,",
    "    merge_com_contexto, atualizar_sessao_pos_turno,",
    ")",
])
p(doc, "📖 Importa o que vamos usar. 'sys' é da biblioteca padrão. As outras são do nosso projeto. "
       "Note que 'from bot.contexto import (...)' importa várias funções de uma vez — os parênteses permitem quebrar em várias linhas.")

py(doc, [
    "def main():",
    "    try:",
    "        sys.stdout.reconfigure(encoding='utf-8')",
    "    except Exception:",
    "        pass",
])
p(doc, "📖 Define a função principal. As 3 linhas dentro do try forçam o terminal a usar UTF-8 (pra acentos "
       "funcionarem no Windows). Se der erro, ignora (pass) — em outros SO não precisa.")

py(doc, [
    "    print('=' * 60)",
    "    print('  Fashion Flow Bot — Atendimento de Produção')",
    "    print('=' * 60)",
])
p(doc, "📖 '=' * 60 cria uma string com 60 sinais de igual (Python permite multiplicar string por número). "
       "Apenas decoração no terminal.")

py(doc, [
    "    dados = carregar_dados()",
    "    sessao = criar_sessao()",
])
p(doc, "📖 Chama as duas funções pra deixar tudo preparado: carrega os CSVs e cria uma sessão nova vazia.")

py(doc, ["    while True:"])
p(doc, "📖 Loop INFINITO — só sai com 'break' ou Ctrl+C.")

py(doc, [
    "        try:",
    "            mensagem = input('Você: ').strip()",
    "        except (EOFError, KeyboardInterrupt):",
    "            print('\\nAté logo!')",
    "            break",
])
p(doc, "📖 input() lê uma linha do usuário. .strip() tira espaços nas pontas. Se o usuário apertar Ctrl+C "
       "(KeyboardInterrupt) ou fechar a entrada (EOFError), captura o erro e quebra o loop com 'break'.")

py(doc, [
    "        if not mensagem:",
    "            continue",
])
p(doc, "📖 Se a mensagem está vazia (só apertou Enter), 'continue' pula pro próximo ciclo do loop sem fazer nada. "
       "'not mensagem' é True quando mensagem é string vazia ''.")

py(doc, [
    "        if mensagem.lower() in ('sair', 'exit', 'quit'):",
    "            print('Bot: Até logo!')",
    "            break",
])
p(doc, "📖 .lower() converte tudo pra minúsculo (assim 'SAIR' e 'Sair' funcionam). "
       "Operador 'in' verifica se a string está numa tupla — equivale a "
       "'mensagem.lower() == \"sair\" OR mensagem.lower() == \"exit\" OR ...'.")

py(doc, [
    "        if mensagem.startswith('/contexto'):",
    "            print(f'  foco_atual: {sessao[\"foco_atual\"]}')",
    "            print(f'  ultimo_assunto: {sessao[\"ultimo_assunto\"]}')",
    "            for i, t in enumerate(sessao['historico_turnos'][-5:], 1):",
    "                print(f'    {i}. [{t[\"intencao\"]}] {t[\"msg\"]!r}')",
    "            continue",
])
p(doc, "📖 Comando de debug. .startswith() checa se a mensagem começa com '/contexto'. "
       "enumerate(...) dá um índice (1, 2, 3...) junto com cada item. "
       "[-5:] pega os 5 últimos. !r dentro do f-string formata com aspas (chama repr())." )

py(doc, [
    "        if is_despedida(mensagem):",
    "            print('Bot: Até logo!')",
    "            sessao = resetar_sessao(sessao)",
    "            continue",
])
p(doc, "📖 Chama nossa função is_despedida. Se for True, despede e ZERA a sessão. "
       "Importante: a função is_despedida foi reescrita pra ser restritiva (só dispara em despedidas puras).")

py(doc, [
    "        if is_casual(mensagem) and sessao['ativa']:",
    "            print('Bot: Beleza, pode continuar!')",
    "            continue",
])
p(doc, "📖 'and' é o operador E lógico — só entra no if se as DUAS condições forem verdadeiras. "
       "Aqui: 'mensagem é casual' E 'sessão está ativa'.")

py(doc, [
    "        em_menu = bool(sessao.get('aguardando_opcao'))",
    "        slots_turno = extrair_slots(mensagem, em_menu=em_menu)",
    "        slots_efetivos = merge_com_contexto(slots_turno, sessao)",
    "        intencao = classificar(mensagem, slots_turno, slots_efetivos, dados['intencoes'], sessao)",
    "        resposta = responder(intencao, slots_efetivos, dados, sessao, mensagem)",
    "        atualizar_sessao_pos_turno(sessao, mensagem, slots_efetivos, intencao, resposta)",
])
p(doc, "📖 O PIPELINE de 5 etapas. Cada linha é uma etapa:")
p(doc, "  1. em_menu = True/False dependendo se há menu aberto.")
p(doc, "  2. Extrai slots da mensagem atual (função pura).")
p(doc, "  3. Mescla com o contexto, aplicando invalidação.")
p(doc, "  4. Classifica a intenção.")
p(doc, "  5. Gera a resposta.")
p(doc, "  6. Persiste mudanças (histórico + foco).")

py(doc, [
    "        print(f'Bot: {resposta}\\n')",
    "",
    "if __name__ == '__main__':",
    "    main()",
])
p(doc, "📖 Imprime a resposta com quebra de linha extra (\\n). "
       "A última estrutura ('if __name__ == ...') é uma convenção do Python: "
       "diz 'só rode main() se esse arquivo for executado diretamente, "
       "não se for importado por outro arquivo'.")

# ── contexto.py — o coração ──
doc.add_page_break()
h(doc, "4.3 bot/contexto.py — a memória da conversa", level=2)
p(doc, "Esse é o arquivo MAIS IMPORTANTE do projeto. Vou explicar linha por linha as funções "
       "mais essenciais.")

py(doc, [
    "SLOTS_EFEMEROS = {'numero_pedido', 'metragem', 'prazo_desejado', 'urgente'}",
])
p(doc, "📖 Cria um SET (parecido com lista, mas sem ordem e sem duplicatas). "
       "Marca quais slots são 'descartáveis' — só valem pro turno em que apareceram. "
       "Set usa { } só com valores (sem 'chave: valor', senão vira dict).")

py(doc, [
    "SLOTS_FILHOS_PRODUTO = {'personalizacao', 'cor', 'grade', 'tecido'}",
])
p(doc, "📖 Quais slots são 'propriedades do produto'. Quando produto muda, esses zeram.")

py(doc, [
    "def criar_sessao():",
    "    return {",
    "        'foco_atual': {},",
    "        'historico_turnos': [],",
    "        'aguardando_opcao': None,",
    "        'ultimo_assunto': None,",
    "        'ativa': False,",
    "    }",
])
p(doc, "📖 Função que cria uma sessão zerada. Retorna um dicionário com 5 campos. "
       "Note: {} é dict vazio, [] é lista vazia, None é o 'nada' do Python.")

py(doc, [
    "def merge_com_contexto(slots_do_turno, sessao):",
    "    foco = dict(sessao.get('foco_atual', {}))",
])
p(doc, "📖 dict(...) cria uma CÓPIA do foco_atual. Se mexer em 'foco' não muda o original — "
       "ainda. (Vamos mudar de propósito mais embaixo, mas controlado.)")

py(doc, [
    "    produto_novo = slots_do_turno.get('produto')",
    "    if produto_novo and foco.get('produto') and produto_novo != foco['produto']:",
    "        for k in SLOTS_FILHOS_PRODUTO:",
    "            foco.pop(k, None)",
])
p(doc, "📖 A REGRA DE INVALIDAÇÃO. Lê passo a passo:")
p(doc, "  • produto_novo = produto do turno atual (pode ser None).")
p(doc, "  • Se o turno trouxe produto E o foco já tinha produto E são DIFERENTES → muda de assunto.")
p(doc, "  • Pra cada 'slot filho' (personalizacao/cor/grade/tecido), remove do foco. "
       ".pop(k, None) tira a chave k; o segundo argumento é o default se a chave não existir (evita erro).")

py(doc, [
    "    for chave, valor in slots_do_turno.items():",
    "        foco[chave] = valor",
    "    return foco",
])
p(doc, "📖 Depois de invalidar, aplica os slots do turno por cima. Devolve o foco resultante.")

nota(doc, "Essa função é o CONSERTO do bug 1 da auditoria. Antes ela não tinha a regra de invalidação — "
          "o produto mudava mas a personalização ficava grudada do produto anterior. Bot mentia.")

py(doc, [
    "def atualizar_sessao_pos_turno(sessao, mensagem, slots_efetivos, intencao, resposta):",
    "    sessao['historico_turnos'].append({",
    "        'msg': mensagem,",
    "        'intencao': intencao,",
    "        'slots': dict(slots_efetivos),",
    "        'resposta': resposta[:200],",
    "    })",
])
p(doc, "📖 Adiciona um registro ao histórico. .append() adiciona no fim de uma lista. "
       "resposta[:200] pega os primeiros 200 caracteres (slice de string).")

py(doc, [
    "    novo_foco = {k: v for k, v in slots_efetivos.items() if k not in SLOTS_EFEMEROS}",
])
p(doc, "📖 DICT COMPREHENSION: cria um novo dict com TODOS os pares (k, v) do slots_efetivos, "
       "EXCETO os que estão na lista de efêmeros. Equivale a:")
py(doc, [
    "novo_foco = {}",
    "for k, v in slots_efetivos.items():",
    "    if k not in SLOTS_EFEMEROS:",
    "        novo_foco[k] = v",
])

py(doc, [
    "    if intencao in INTENCOES_PEDIDO and sessao.get('ultimo_assunto') in INTENCOES_ORCAMENTO:",
    "        novo_foco.pop('produto', None)",
    "        novo_foco.pop('quantidade', None)",
])
p(doc, "📖 Outra regra de invalidação: se o usuário mudou de 'orçamento' (preço/prazo) pra 'pedido' "
       "(status/cancelar), zera produto/quantidade — não fazem sentido juntos.")

py(doc, [
    "    sessao['foco_atual'] = novo_foco",
    "    if intencao != 'selecao_opcao':",
    "        sessao['ultimo_assunto'] = intencao",
    "    sessao['ativa'] = True",
    "    return sessao",
])
p(doc, "📖 Atualiza o foco. ultimo_assunto NÃO atualiza se a intenção foi 'selecao_opcao' "
       "(o usuário só respondeu a um menu — o assunto real continua sendo o anterior).")

py(doc, [
    "def is_despedida(mensagem):",
    "    t = normalizar(mensagem).strip()",
    "    if t in _DESPEDIDAS_EXATAS:",
    "        return True",
    "    palavras = re.findall(r'\\w+', t)",
    "    if 1 <= len(palavras) <= 3:",
    "        if all(p in _DESPEDIDAS_EXATAS or p in _CASUAIS for p in palavras):",
    "            if any(p in _DESPEDIDAS_EXATAS for p in palavras):",
    "                return True",
    "    return False",
])
p(doc, "📖 Função restritiva de despedida. re.findall(r'\\w+', t) pega todas as 'palavras' (sequências de letras/dígitos) da string. "
       "all(...) é True se TODAS as condições forem verdadeiras; any(...) se PELO MENOS UMA for. "
       "Condições combinadas: mensagem curta (1-3 palavras) E todas as palavras são despedida/casual E pelo menos uma é despedida pura.")

# ── extractor.py — partes essenciais ──
doc.add_page_break()
h(doc, "4.4 bot/extractor.py — partes essenciais", level=2)
p(doc, "Arquivo grande. Vou explicar os trechos mais didáticos e que aparecem mais em prova.")

py(doc, [
    "def normalizar(texto):",
    "    texto = texto.lower()",
    "    texto = unicodedata.normalize('NFD', texto)",
    "    texto = ''.join(c for c in texto if unicodedata.category(c) != 'Mn')",
    "    return texto",
])
p(doc, "📖 Tira acentos e converte pra minúsculo. unicodedata.normalize('NFD', ...) decompõe "
       "letras acentuadas em letra base + acento separado. Depois ''.join(...) reúne apenas as letras base "
       "(unicodedata.category(c) != 'Mn' filtra os acentos — Mn = 'Mark, nonspacing').")

py(doc, [
    "def extrair_slots(mensagem, em_menu=False):",
    "    if em_menu:",
    "        return {}",
    "    t = normalizar(mensagem)",
    "    slots = {}",
])
p(doc, "📖 Se está no meio de um menu, devolve dict vazio (evita confundir '1' do menu com quantidade).")

py(doc, [
    "    match_pedido = re.search(r'FF-\\d{4}-\\d{4}', mensagem, re.IGNORECASE)",
    "    if match_pedido:",
    "        slots['numero_pedido'] = match_pedido.group(0).upper()",
    "        t = re.sub(r'ff-\\d{4}-\\d{4}', ' ', t)",
])
p(doc, "📖 Procura número de pedido. re.IGNORECASE faz busca case-insensitive. "
       "match.group(0) é o texto inteiro que casou. "
       "Depois re.sub(...) REMOVE o número de pedido do texto, substituindo por espaço — "
       "isso evita que '2024' do código vire 'quantidade'. CONSERTO do bug 2.")

py(doc, [
    "    match_qtd = re.search(rf'\\b(\\d+)\\s*({UNIDADES_QTD})\\b', t)",
    "    if match_qtd:",
    "        slots['quantidade'] = int(match_qtd.group(1))",
])
p(doc, "📖 rf'...' é uma f-string COM regex (r). UNIDADES_QTD é uma constante com lista de unidades aceitas "
       "('pecas?|camisetas?|...'). match.group(1) pega o que está no primeiro parênteses — o número. "
       "int(...) converte de string '100' pra inteiro 100.")

py(doc, [
    "    neg_pers = re.search(",
    "        r'(?:sem|nao quero|sem nenhuma?)\\s+'",
    "        r'(bordad[oa]s?|silk\\w*|serigrafia|dtf|estamp\\w*|etiqueta|personali\\w+)',",
    "        t",
    "    )",
    "    if neg_pers:",
    "        slots['personalizacao'] = 'nenhuma'",
])
p(doc, "📖 Detecta negação. (?:...) é um grupo SEM captura (só agrupa, não vira group(1)). "
       "\\s+ é um ou mais espaços. Se acha 'sem bordado', 'sem silk', etc, marca personalizacao='nenhuma'. "
       "CONSERTO do bug 22.")

# ── classifier.py ──
h(doc, "4.5 bot/classifier.py — esqueleto", level=2)
p(doc, "Aqui vou explicar a ESTRUTURA, não cada regra. As regras são longas, mas o padrão se repete.")

py(doc, [
    "def classificar(mensagem, slots_turno, slots_efetivos, intencoes, sessao=None):",
    "    t = normalizar(mensagem)",
    "",
    "    # 1. Menu ativo?",
    "    if sessao and sessao.get('aguardando_opcao'):",
    "        return 'selecao_opcao'",
    "",
    "    # 2. Verbos (cancelar, alterar, estoque)",
    "    if re.search(r'\\bcancelar\\b', t):",
    "        return 'cancelar_pedido'",
    "    # ...",
    "",
    "    # 3. Slots do turno (não da sessão)",
    "    if slots_turno.get('numero_pedido'):",
    "        return 'status_pedido'",
    "",
    "    # 4. Combinações (preço/prazo)",
    "    if quantidade and produto:",
    "        if re.search(r'custa|preco', t):",
    "            return 'combinado_preco_qtd_produto'",
    "    # ...",
    "",
    "    # 5. Fallback",
    "    return 'fallback'",
])
p(doc, "📖 PADRÃO: várias verificações 'if ... return'. A primeira que casar VENCE. Por isso a ORDEM "
       "das regras é estratégica — verbo (cancelar) antes de slot (numero_pedido), senão 'cancelar pedido X' "
       "vira 'status do pedido X'.")

# ── responder.py — estrutura ──
h(doc, "4.6 bot/responder.py — padrão dos handlers", level=2)
p(doc, "Esse arquivo é o maior, mas todos os blocos seguem o mesmo padrão. Vou mostrar 1 bloco completo.")

py(doc, [
    "if intencao == 'combinado_preco_qtd_produto':",
    "    quantidade = slots.get('quantidade', 0)",
    "    produto = slots.get('produto')",
    "    personalizacao = slots.get('personalizacao', 'nenhuma')",
])
p(doc, "📖 .get(chave, default) pega o valor ou usa o default se a chave não existe. Aqui: "
       "quantidade vira 0 se não tem, personalizacao vira 'nenhuma'.")

py(doc, [
    "    if not produto:",
    "        return (",
    "            'Pra calcular preço eu preciso saber qual produto você quer. '",
    "            'Trabalhamos com camisetas, polos, moletons... Qual deles?'",
    "        )",
])
p(doc, "📖 Pergunta clarificadora quando falta info. Note: 2 strings entre parênteses ficam concatenadas "
       "automaticamente — é forma comum de quebrar strings longas em várias linhas.")

py(doc, [
    "    df = dados['preco']",
    "    filtro = df[",
    "        (df['produto'] == produto) &",
    "        (df['qtd_min'] <= quantidade) &",
    "        (df['qtd_max'] >= quantidade) &",
    "        (df['personalizacao'] == personalizacao)",
    "    ]",
])
p(doc, "📖 Filtra o DataFrame com 4 condições combinadas por & (and em Pandas usa & com parênteses). "
       "df[booleano] devolve só as linhas onde a condição é True.")

py(doc, [
    "    if filtro.empty:",
    "        return f'Não encontrei dados de preço para {produto} com quantidade {quantidade}.'",
])
p(doc, "📖 .empty é True se o filtro não retornou nenhuma linha. Caso não encontre, mensagem amigável.")

py(doc, [
    "    r = filtro.iloc[0]",
    "    total = round(float(r['preco_unitario_estimado']) * quantidade, 2)",
    "    pers_txt = f' com {personalizacao}' if personalizacao != 'nenhuma' else ''",
    "    return (",
    "        f'Para {quantidade} {pecas(quantidade)} de {produto.replace(\"_\", \" \")}{pers_txt}: '",
    "        f'valor unitário de R$ {r[\"preco_unitario_estimado\"]}. '",
    "        f'Total estimado: R$ {total:.2f}.'",
    "    )",
])
p(doc, "📖 .iloc[0] pega a primeira linha do filtro. round(x, 2) arredonda pra 2 casas. "
       "OPERADOR TERNÁRIO: 'X if condição else Y' — devolve X se a condição é verdadeira, senão Y. "
       "pecas() é nossa função helper pra singular/plural.")

# ── app.py ──
h(doc, "4.7 app.py — API REST com FastAPI", level=2)

py(doc, [
    "from fastapi import FastAPI",
    "from fastapi.middleware.cors import CORSMiddleware",
    "from fastapi.responses import FileResponse",
    "from pydantic import BaseModel",
    "",
    "app = FastAPI(title='Fashion Flow Bot')",
])
p(doc, "📖 Cria a aplicação. FastAPI usa Pydantic pra validar dados de entrada (mais sobre isso já já).")

py(doc, [
    "app.add_middleware(",
    "    CORSMiddleware,",
    "    allow_origins=['*'],",
    "    allow_methods=['*'],",
    "    allow_headers=['*'],",
    ")",
])
p(doc, "📖 CORS permite que páginas em outro domínio chamem nossa API. allow_origins=['*'] significa "
       "'qualquer site'. Em produção real, restrinja pra domínios específicos.")

py(doc, [
    "dados = carregar_dados()",
    "sessoes = {}",
])
p(doc, "📖 Carrega os dados uma vez no startup (não em cada request). 'sessoes' é dict global "
       "onde guardamos cada sessão pelo seu ID.")

py(doc, [
    "class MensagemRequest(BaseModel):",
    "    sessao_id: str",
    "    mensagem: str",
])
p(doc, "📖 Define uma 'classe' que descreve o formato esperado do JSON. FastAPI valida automaticamente: "
       "se faltar campo ou tipo errado, retorna erro 422.")

py(doc, [
    "@app.post('/chat')",
    "def chat(req: MensagemRequest):",
    "    sessao_id = req.sessao_id",
    "    mensagem = req.mensagem.strip()",
])
p(doc, "📖 Endpoint POST /chat. O parâmetro req é validado como MensagemRequest automaticamente.")

py(doc, [
    "    if sessao_id not in sessoes:",
    "        sessoes[sessao_id] = criar_sessao()",
    "    sessao = sessoes[sessao_id]",
])
p(doc, "📖 Se é a primeira vez desse usuário, cria sessão nova. Senão recupera a existente.")

py(doc, [
    "    # pipeline igual ao do main.py",
    "    em_menu = bool(sessao.get('aguardando_opcao'))",
    "    slots_turno = extrair_slots(mensagem, em_menu=em_menu)",
    "    slots_efetivos = merge_com_contexto(slots_turno, sessao)",
    "    intencao = classificar(mensagem, slots_turno, slots_efetivos, dados['intencoes'], sessao)",
    "    resposta = responder(intencao, slots_efetivos, dados, sessao, mensagem)",
    "    atualizar_sessao_pos_turno(sessao, mensagem, slots_efetivos, intencao, resposta)",
    "",
    "    return {'resposta': resposta, 'intencao': intencao}",
])
p(doc, "📖 Mesmo pipeline do main.py. Retorna um dict — FastAPI converte automaticamente em JSON.")

doc.add_page_break()


# ##################################################################
# PARTE V — SIMULADO
# ##################################################################
h(doc, "Parte V — Simulado: questões prováveis", level=1, color=ROXO)
p(doc, "Questões no estilo do que sua professora pode pedir. Tenta resolver SOZINHA primeiro, "
       "depois confere o gabarito. Cobrimos os 6 tópicos clássicos de 1º sem: "
       "entrada/saída, condicional, loop, função, lista/vetor e dicionário/registro.")

# Q1
h(doc, "Questão 1 — Entrada e condicional", level=2, color=AZUL)
p(doc, "Escreva um programa que peça uma idade ao usuário e classifique:")
doc.add_paragraph("Menor de 18: 'menor de idade'", style="List Bullet")
doc.add_paragraph("Entre 18 e 64: 'adulto'", style="List Bullet")
doc.add_paragraph("65 ou mais: 'idoso'", style="List Bullet")

add_run(p(doc), "Gabarito Python:", bold=True)
py(doc, [
    "idade = int(input('Digite sua idade: '))",
    "",
    "if idade < 18:",
    "    print('menor de idade')",
    "elif idade < 65:",
    "    print('adulto')",
    "else:",
    "    print('idoso')",
])

add_run(p(doc), "Gabarito Portugol:", bold=True)
portugol(doc, [
    'algoritmo "classifica_idade"',
    "var",
    "   idade: inteiro",
    "inicio",
    '   escreva("Digite sua idade: ")',
    "   leia(idade)",
    "   se idade < 18 entao",
    '      escreva("menor de idade")',
    "   senao se idade < 65 entao",
    '      escreva("adulto")',
    "   senao",
    '      escreva("idoso")',
    "   fimse",
    "fimalgoritmo",
])

# Q2
h(doc, "Questão 2 — Loop com contador", level=2, color=AZUL)
p(doc, "Escreva um programa que mostre os números pares de 2 a 20.")

add_run(p(doc), "Gabarito Python (com range):", bold=True)
py(doc, [
    "for n in range(2, 21, 2):",
    "    print(n)",
])

add_run(p(doc), "Gabarito Python (com if):", bold=True)
py(doc, [
    "for n in range(1, 21):",
    "    if n % 2 == 0:",
    "        print(n)",
])

add_run(p(doc), "Gabarito Portugol:", bold=True)
portugol(doc, [
    "para n de 2 ate 20 passo 2 faca",
    "   escreva(n)",
    "fimpara",
])

# Q3
h(doc, "Questão 3 — Função (média)", level=2, color=AZUL)
p(doc, "Escreva uma função que receba 3 notas e retorne a média aritmética. "
       "Em seguida, peça as notas ao usuário e mostre a média.")

add_run(p(doc), "Gabarito Python:", bold=True)
py(doc, [
    "def media(n1, n2, n3):",
    "    return (n1 + n2 + n3) / 3",
    "",
    "a = float(input('Nota 1: '))",
    "b = float(input('Nota 2: '))",
    "c = float(input('Nota 3: '))",
    "",
    "m = media(a, b, c)",
    "print(f'Média: {m:.2f}')",
])

add_run(p(doc), "Gabarito Portugol:", bold=True)
portugol(doc, [
    'algoritmo "media"',
    "var",
    "   a, b, c, m: real",
    "",
    "funcao calcular_media(x: real, y: real, z: real): real",
    "inicio",
    "   retorne (x + y + z) / 3",
    "fimfuncao",
    "",
    "inicio",
    '   escreva("Nota 1: ")',
    "   leia(a)",
    '   escreva("Nota 2: ")',
    "   leia(b)",
    '   escreva("Nota 3: ")',
    "   leia(c)",
    "   m <- calcular_media(a, b, c)",
    '   escreva("Média: ", m)',
    "fimalgoritmo",
])

# Q4
h(doc, "Questão 4 — Lista/vetor (maior valor)", level=2, color=AZUL)
p(doc, "Leia 5 números e mostre qual é o maior.")

add_run(p(doc), "Gabarito Python:", bold=True)
py(doc, [
    "numeros = []",
    "for i in range(5):",
    "    n = float(input(f'Número {i+1}: '))",
    "    numeros.append(n)",
    "",
    "maior = numeros[0]",
    "for n in numeros:",
    "    if n > maior:",
    "        maior = n",
    "",
    "print(f'Maior: {maior}')",
    "",
    "# Versão curtinha (usando função pronta):",
    "print(f'Maior: {max(numeros)}')",
])

add_run(p(doc), "Gabarito Portugol:", bold=True)
portugol(doc, [
    "var",
    "   numeros: vetor[1..5] de real",
    "   i: inteiro",
    "   maior: real",
    "inicio",
    "   para i de 1 ate 5 faca",
    '      escreva("Numero ", i, ": ")',
    "      leia(numeros[i])",
    "   fimpara",
    "",
    "   maior <- numeros[1]",
    "   para i de 2 ate 5 faca",
    "      se numeros[i] > maior entao",
    "         maior <- numeros[i]",
    "      fimse",
    "   fimpara",
    "",
    '   escreva("Maior: ", maior)',
])

# Q5
h(doc, "Questão 5 — Dicionário (estoque simples)", level=2, color=AZUL)
p(doc, "Crie um dicionário com produtos e quantidades em estoque. Peça ao usuário o nome de "
       "um produto e mostre quantos itens tem (ou 'produto não encontrado').")

add_run(p(doc), "Gabarito Python:", bold=True)
py(doc, [
    "estoque = {",
    "    'camiseta': 50,",
    "    'polo': 30,",
    "    'moletom': 12,",
    "    'jeans': 0,",
    "}",
    "",
    "produto = input('Qual produto? ').lower()",
    "",
    "if produto in estoque:",
    "    qtd = estoque[produto]",
    "    if qtd == 0:",
    "        print(f'{produto}: sem estoque')",
    "    else:",
    "        print(f'{produto}: {qtd} unidades')",
    "else:",
    "    print('Produto não encontrado')",
])

p(doc, "Em portugol não tem dicionário nativo, então normalmente usamos vetores paralelos:")
portugol(doc, [
    "var",
    "   nomes: vetor[1..4] de caractere",
    "   qtds:  vetor[1..4] de inteiro",
    "   busca: caractere",
    "   i: inteiro",
    "   achou: logico",
    "inicio",
    '   nomes[1] <- "camiseta"; qtds[1] <- 50',
    '   nomes[2] <- "polo";     qtds[2] <- 30',
    '   nomes[3] <- "moletom";  qtds[3] <- 12',
    '   nomes[4] <- "jeans";    qtds[4] <- 0',
    "",
    '   escreva("Qual produto? ")',
    "   leia(busca)",
    "",
    "   achou <- falso",
    "   para i de 1 ate 4 faca",
    "      se nomes[i] = busca entao",
    '         escreva(busca, ": ", qtds[i], " unidades")',
    "         achou <- verdadeiro",
    "      fimse",
    "   fimpara",
    "",
    "   se nao achou entao",
    '      escreva("Produto não encontrado")',
    "   fimse",
])

# Q6
h(doc, "Questão 6 — Tudo junto (mini bot)", level=2, color=AZUL)
p(doc, "Escreva um mini chatbot que: pergunta o nome do usuário, depois fica em loop "
       "respondendo perguntas até o usuário digitar 'sair'. Reconhece 3 'intenções': "
       "'oi/olá' (saudação), 'preço' (informa preço fictício), 'sair' (encerra).")

add_run(p(doc), "Gabarito Python:", bold=True)
py(doc, [
    "def classificar(msg):",
    "    msg = msg.lower()",
    "    if 'oi' in msg or 'olá' in msg or 'ola' in msg:",
    "        return 'saudacao'",
    "    if 'preço' in msg or 'preco' in msg or 'custa' in msg:",
    "        return 'preco'",
    "    if msg == 'sair':",
    "        return 'sair'",
    "    return 'desconhecido'",
    "",
    "def responder(intencao, nome):",
    "    if intencao == 'saudacao':",
    "        return f'Olá, {nome}! Como posso ajudar?'",
    "    if intencao == 'preco':",
    "        return 'A camiseta custa R$ 29,90.'",
    "    if intencao == 'sair':",
    "        return f'Tchau, {nome}!'",
    "    return 'Desculpa, não entendi.'",
    "",
    "nome = input('Qual seu nome? ')",
    "print(f'Oi, {nome}! Pode perguntar (ou digite sair).')",
    "",
    "while True:",
    "    msg = input('Você: ')",
    "    intencao = classificar(msg)",
    "    resp = responder(intencao, nome)",
    "    print(f'Bot: {resp}')",
    "    if intencao == 'sair':",
    "        break",
])

nota(doc, "Esse último exemplo é literalmente o esqueleto do nosso projeto, em miniatura! "
          "É uma boa chance pra você ver TODOS os conceitos juntos: input, while, função, if, dict, return.")

doc.add_page_break()

# ── GLOSSÁRIO FINAL ──
h(doc, "Glossário rápido", level=1, color=ROXO)

tab = doc.add_table(rows=20, cols=2)
tab.style = "Light Grid Accent 1"
hdr = ["Termo", "Significado"]
for j, lbl in enumerate(hdr):
    tab.rows[0].cells[j].text = ""
    add_run(tab.rows[0].cells[j].paragraphs[0], lbl, bold=True)

termos = [
    ("Variável", "Um nome que aponta pra um valor (uma 'caixa' nomeada)"),
    ("Tipo", "Que espécie de dado a variável guarda (int, str, bool...)"),
    ("Atribuição", "Guardar um valor numa variável (x = 10)"),
    ("Operador", "Símbolo que faz operação (+, -, ==, and, or, not)"),
    ("Função", "Pedaço de código com nome, que pode ser reusado"),
    ("Parâmetro", "Variável que a função recebe quando é chamada"),
    ("Return", "Valor que a função devolve pra quem chamou"),
    ("Lista", "Sequência ordenada de valores (uses [ ])"),
    ("Dicionário", "Coleção de chave → valor (uses { 'chave': valor })"),
    ("Set", "Coleção sem ordem nem duplicatas (uses { valor1, valor2 })"),
    ("None", "O 'nada' do Python — valor especial pra 'sem valor'"),
    ("Indentação", "Espaços antes da linha — definem o que está dentro de quê"),
    ("Loop", "Repetição (for, while)"),
    ("Iterar", "Percorrer item por item de uma sequência"),
    ("Import", "Trazer código de outro arquivo ou biblioteca"),
    ("Módulo", "Um arquivo .py que pode ser importado"),
    ("Regex", "Padrão pra buscar texto"),
    ("DataFrame", "Tabela do pandas (tipo Excel dentro do Python)"),
    ("Side effect", "Quando uma função mexe em algo além de seu retorno"),
]
for i, (k, v) in enumerate(termos, 1):
    tab.rows[i].cells[0].text = k
    tab.rows[i].cells[1].text = v

p(doc)
fim = p(doc)
fim.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(fim, "— fim do material de estudo. Boa prova! 🍀 —",
        italic=True, color=CINZA, size=12)

doc.save(OUT)
print(f"OK: {OUT}")
print(f"Tamanho: {OUT.stat().st_size} bytes")
